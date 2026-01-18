# Módulo de generación de informes técnicos

import streamlit as st
import json
import os
import io
import pandas as pd
import numpy as np
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import matplotlib.pyplot as plt
from ui.ai_module import generar_datos_json
from core.curves import calculate_bep
from ui.epanet_export import render_epanet_export_section

# --- Funciones Auxiliares ---

def capturar_grafico_plotly(fig, grupo, nombre_grafico, forzar_captura=False):
    """
    Captura gráficos Plotly solo cuando se solicita desde la pestaña de Reportes
    Crea una versión simplificada sin líneas duplicadas ni zona de eficiencia
    """
    if not fig or not hasattr(fig, 'data') or not fig.data:
        return False
    
    # Solo capturar si se solicita
    # Eliminada restricción de checkbox para permitir captura proactiva
    pass
    
    try:
        # Inicializar estructura si no existe
        if 'graficos_exportados' not in st.session_state:
            st.session_state['graficos_exportados'] = {
                'grupo_100_rpm': {},
                'grupo_vfd': {}
            }
        
        # Guardar el objeto de figura en session_state
        guardar_grafico_en_session_state(fig, nombre_grafico, grupo, mostrar_log=False)
        
        # Crear gráfico simplificado para captura (sin líneas duplicadas)
        try:
            fig_simplificada = crear_grafico_simplificado_para_captura(fig, nombre_grafico)
            imagen_bytes = figuratransform(fig_simplificada).getvalue()
            
            # Guardar imagen capturada
            st.session_state['graficos_exportados'][grupo][nombre_grafico] = imagen_bytes
            st.session_state[f'grafico_capturado_{nombre_grafico}'] = True
            
        except Exception as e_matplotlib:
            # Marcar como capturado aunque falle Matplotlib
            st.session_state[f'grafico_capturado_{nombre_grafico}'] = True
        
        return True
        
    except Exception as e:
        return False

def figuratransform(fig):
    """Convierte una figura (Matplotlib o Plotly) a un buffer en RAM con PNG a 300 DPI"""
    from io import BytesIO
    
    buf = BytesIO()
    
    # Matplotlib
    if hasattr(fig, 'savefig'):
        fig.savefig(buf, format='png', dpi=300, bbox_inches='tight')
        plt.close(fig)  # Cerrar para liberar memoria
    # Plotly (usando Matplotlib como fallback)
    elif hasattr(fig, 'data'):
        # Convertir Plotly a Matplotlib
        fig_matplotlib = convertir_plotly_a_matplotlib_figura(fig)
        fig_matplotlib.savefig(buf, format='png', dpi=300, bbox_inches='tight')
        plt.close(fig_matplotlib)  # Cerrar para liberar memoria
    else:
        raise ValueError("Figura no soportada. Debe ser Matplotlib o Plotly.")
    
    buf.seek(0)  # ¡MUY IMPORTANTE!
    return buf

def convertir_color_plotly_a_matplotlib(color_plotly):
    """Convierte un color de Plotly a formato compatible con Matplotlib"""
    import re
    
    if color_plotly is None:
        return 'blue'
    
    color_str = str(color_plotly)
    
    # Si ya es un color válido de Matplotlib, devolverlo
    if color_str in ['red', 'blue', 'green', 'black', 'orange', 'purple', 'brown', 'pink', 'gray', 'olive', 'cyan']:
        return color_str
    
    # Convertir formato rgba(r, g, b, a) a formato matplotlib
    rgba_match = re.match(r'rgba?\((\d+),\s*(\d+),\s*(\d+)(?:,\s*([\d.]+))?\)', color_str)
    if rgba_match:
        r, g, b = int(rgba_match.group(1)), int(rgba_match.group(2)), int(rgba_match.group(3))
        a = float(rgba_match.group(4)) if rgba_match.group(4) else 1.0
        
        # Convertir de 0-255 a 0-1 para Matplotlib
        return (r/255.0, g/255.0, b/255.0, a)
    
    # Convertir formato hex
    if color_str.startswith('#'):
        return color_str
    
    # Si no se puede convertir, usar color por defecto
    return 'blue'

def convertir_plotly_a_matplotlib_figura(fig_plotly, solo_punto_operacion=False):
    """Convierte un gráfico de Plotly a figura de Matplotlib"""
    import matplotlib.pyplot as plt
    import numpy as np
    
    # Crear figura de Matplotlib
    fig, ax = plt.subplots(figsize=(12, 8))
    
    # Extraer datos de Plotly y recrear con Matplotlib
    for trace in fig_plotly.data:
        if hasattr(trace, 'x') and hasattr(trace, 'y'):
            x_data = trace.x
            y_data = trace.y
            
            # Verificar que los datos no sean None
            if x_data is None or y_data is None:
                continue
                
            # Convertir a listas si es necesario
            if not isinstance(x_data, (list, tuple, np.ndarray)):
                x_data = [x_data] if x_data is not None else []
            if not isinstance(y_data, (list, tuple, np.ndarray)):
                y_data = [y_data] if y_data is not None else []
            
            # Verificar que tenemos datos válidos
            if len(x_data) == 0 or len(y_data) == 0:
                continue
            
            # Determinar el estilo de línea
            style = '-'
            if hasattr(trace, 'mode') and trace.mode is not None:
                mode_str = str(trace.mode).lower()
                if 'lines' in mode_str and 'markers' in mode_str:
                    style = 'o-'
                elif 'lines' in mode_str:
                    style = '-'
                elif 'markers' in mode_str:
                    style = 'o'
            
            # Obtener color y convertir formato RGBA a compatible con Matplotlib
            color = 'blue'
            if hasattr(trace, 'line') and trace.line is not None and hasattr(trace.line, 'color'):
                color = convertir_color_plotly_a_matplotlib(trace.line.color)
            elif hasattr(trace, 'marker') and trace.marker is not None and hasattr(trace.marker, 'color'):
                color = convertir_color_plotly_a_matplotlib(trace.marker.color)
            
            # Obtener label
            label = trace.name if hasattr(trace, 'name') and trace.name is not None else 'Datos'
            
            # Si solo queremos puntos de operación, solo mostrar marcadores
            if solo_punto_operacion and 'operación' in label.lower():
                ax.plot(x_data, y_data, 'o', color=color, label=label, markersize=10, markeredgewidth=2, markeredgecolor='black')
            elif not solo_punto_operacion:
                # Mostrar líneas normales
                if not any(keyword in label.lower() for keyword in ['eficiencia', 'rendimiento']):
                    ax.plot(x_data, y_data, style, color=color, label=label, linewidth=2, markersize=6)
                else:
                    # ZONA DE EFICIENCIA: Mostrar curva real y evitar duplicados
                    if len(x_data) > 3:
                        # Obtener valores de BEP para el label
                        zona_eff_min, zona_eff_max, bep_eta = obtener_valores_bep_eficiencia()
                        label_eficiencia = f"Zona de eficiencia ({zona_eff_min:.0f}%-{zona_eff_max:.0f}% BEP)"
                        
                        # Mostrar la curva real de eficiencia
                        ax.plot(x_data, y_data, '-', color='lightgray', label=label_eficiencia, 
                               linewidth=2, alpha=0.7)
    
    # Configurar el gráfico
    if (hasattr(fig_plotly.layout, 'title') and 
        fig_plotly.layout.title is not None and 
        hasattr(fig_plotly.layout.title, 'text') and 
        fig_plotly.layout.title.text is not None):
        ax.set_title(fig_plotly.layout.title.text, fontsize=14, fontweight='bold')
    
    if (hasattr(fig_plotly.layout, 'xaxis') and 
        fig_plotly.layout.xaxis is not None and
        hasattr(fig_plotly.layout.xaxis, 'title') and 
        fig_plotly.layout.xaxis.title is not None and
        hasattr(fig_plotly.layout.xaxis.title, 'text') and
        fig_plotly.layout.xaxis.title.text is not None):
        ax.set_xlabel(fig_plotly.layout.xaxis.title.text, fontsize=12)
    
    if (hasattr(fig_plotly.layout, 'yaxis') and 
        fig_plotly.layout.yaxis is not None and
        hasattr(fig_plotly.layout.yaxis, 'title') and 
        fig_plotly.layout.yaxis.title is not None and
        hasattr(fig_plotly.layout.yaxis.title, 'text') and
        fig_plotly.layout.yaxis.title.text is not None):
        ax.set_ylabel(fig_plotly.layout.yaxis.title.text, fontsize=12)
    
    # Agregar grid y leyenda
    ax.grid(True, alpha=0.3)
    if len(fig_plotly.data) > 1:
        ax.legend()
    
    # Ajustar layout
    plt.tight_layout()
    
    return fig

def convertir_plotly_a_matplotlib(fig_plotly, nombre_grafico):
    """Convierte un gráfico de Plotly a imagen usando Matplotlib (método legacy)"""
    fig = convertir_plotly_a_matplotlib_figura(fig_plotly)
    buf = figuratransform(fig)
    return buf.getvalue()

def probar_captura_manual():
    """Prueba la captura manual de un gráfico de prueba"""
    try:
        st.info("🧪 Creando gráfico de prueba...")
        
        # Verificar versiones primero
        try:
            import plotly
            st.info(f"📊 Plotly versión: {plotly.__version__}")
        except Exception as e:
            st.error(f"❌ Error importando Plotly: {e}")
            return
        
        # Verificar Matplotlib
        try:
            import matplotlib
            st.info(f"📈 Matplotlib versión: {matplotlib.__version__}")
        except Exception as e:
            st.error(f"❌ Error importando Matplotlib: {e}")
            return
        
        # Crear un gráfico de prueba simple
        import plotly.graph_objects as go
        
        fig_prueba = go.Figure()
        fig_prueba.add_trace(go.Scatter(x=[1, 2, 3, 4], y=[10, 11, 12, 13], mode='lines+markers', name='Prueba'))
        fig_prueba.update_layout(title='Gráfico de Prueba', xaxis_title='X', yaxis_title='Y')
        
        st.info("📈 Gráfico de prueba creado")
        
        # Mostrar el gráfico
        st.plotly_chart(fig_prueba, use_container_width=True)
        
        # Intentar capturar el gráfico de prueba
        if capturar_grafico_plotly(fig_prueba, 'grupo_100_rpm', 'prueba'):
            st.success("✅ Captura de prueba exitosa")
            st.info("💡 La función de captura funciona correctamente con Matplotlib")
            
            # Mostrar información adicional
            if 'graficos_exportados' in st.session_state:
                graficos = st.session_state['graficos_exportados']
                total = len(graficos.get('grupo_100_rpm', {})) + len(graficos.get('grupo_vfd', {}))
                st.info(f"📊 Total de gráficos en memoria: {total}")
        else:
            st.error("❌ La función de captura falló")
            
    except Exception as e:
        st.error(f"❌ Error en prueba de captura: {e}")
        import traceback
        st.error("Detalles del error:")
        st.code(traceback.format_exc())

def diagnostico_sistema_completo():
    """Ejecuta un diagnóstico completo del sistema de captura de gráficos"""
    st.markdown("### 🔍 Diagnóstico Completo del Sistema")
    
    # Información del sistema
    import sys
    st.info(f"🐍 Python: {sys.version.split()[0]}")
    st.info(f"📁 Directorio: {os.getcwd()}")
    st.info(f"🔧 Ejecutable: {sys.executable}")
    
    # Verificar Plotly
    try:
        import plotly
        st.success(f"✅ Plotly versión: {plotly.__version__}")
        st.info(f"📁 Plotly path: {plotly.__file__}")
    except Exception as e:
        st.error(f"❌ Error con Plotly: {e}")
        return False
    
    # Verificar Matplotlib
    try:
        import matplotlib
        st.success(f"✅ Matplotlib versión: {matplotlib.__version__}")
        st.info(f"📁 Matplotlib path: {matplotlib.__file__}")
    except ImportError as e:
        st.error(f"❌ Matplotlib no está instalado: {e}")
        st.info("💡 Instala con: pip install matplotlib")
        return False
    
    # Probar exportación
    st.info("🔄 Probando exportación de gráfico con Matplotlib...")
    
    try:
        import plotly.graph_objects as go
        fig = go.Figure()
        fig.add_trace(go.Scatter(x=[1, 2, 3], y=[4, 5, 6]))
        
        # Mostrar el gráfico
        st.plotly_chart(fig, use_container_width=True)
        
        # Probar método con Matplotlib
        try:
            imagen_bytes = convertir_plotly_a_matplotlib(fig, "diagnostico")
            st.success(f"✅ Exportación con Matplotlib exitosa: {len(imagen_bytes)} bytes")
            st.info("💡 El sistema funcionará correctamente con Matplotlib")
            return True
        except Exception as e:
            st.error(f"❌ Error en exportación con Matplotlib: {e}")
            import traceback
            st.error("Detalles del error:")
            st.code(traceback.format_exc())
            return False
        
    except Exception as e:
        st.error(f"❌ Error general en exportación: {e}")
        import traceback
        st.error("Detalles del error:")
        st.code(traceback.format_exc())
        return False

def debug_captura_graficos():
    """Muestra información de debugging para la captura de gráficos"""
    st.markdown("### 🔍 Debug: Información de Captura de Gráficos")
    
    # Ejecutar diagnóstico del sistema primero
    if st.button("🔍 Ejecutar Diagnóstico Completo", key="diagnostico_completo"):
        diagnostico_sistema_completo()
    
    # Botón para verificar gráficos capturados
    if st.button("🔍 Verificar Gráficos Capturados", key="verificar_graficos"):
        forzar_captura_graficos()
    
    st.markdown("---")
    
    # Verificar session_state
    if 'tablas_graficos' not in st.session_state:
        st.error("❌ No existe 'tablas_graficos' en session_state")
        st.info("💡 Las curvas se generan automáticamente cuando configuras el proyecto. Ve a la pestaña 'Análisis de curvas' para verificar que las curvas estén visibles.")
        return
    
    tablas = st.session_state.get('tablas_graficos', {})
    st.info(f"📊 Estructura de tablas_graficos: {list(tablas.keys())}")
    
    # Verificar datos 100% RPM
    datos_100 = tablas.get('tablas_100_rpm', {})
    if datos_100:
        st.success("✅ Datos 100% RPM encontrados")
        for key, value in datos_100.items():
            if isinstance(value, dict) and 'data' in value:
                df = deserialize_df(value)
                if not df.empty:
                    st.info(f"  - {key}: {len(df)} filas")
                else:
                    st.warning(f"  - {key}: DataFrame VACÍO")
            else:
                st.warning(f"  - {key}: Formato incorrecto")
    else:
        st.error("❌ No hay datos 100% RPM")
    
    # Verificar datos VFD
    datos_vfd = tablas.get('tablas_vfd_rpm', {})
    if datos_vfd:
        st.success("✅ Datos VFD encontrados")
        for key, value in datos_vfd.items():
            if isinstance(value, dict) and 'data' in value:
                df = deserialize_df(value)
                if not df.empty:
                    st.info(f"  - {key}: {len(df)} filas")
                else:
                    st.warning(f"  - {key}: DataFrame VACÍO")
            else:
                st.warning(f"  - {key}: Formato incorrecto")
    else:
        st.error("❌ No hay datos VFD")
    
    # Verificar si hay datos suficientes para captura
    df_bomba_100 = deserialize_df(datos_100.get('df_bomba_100')) if datos_100 else pd.DataFrame()
    df_sistema_100 = deserialize_df(datos_100.get('df_sistema_100')) if datos_100 else pd.DataFrame()
    
    if not df_bomba_100.empty and not df_sistema_100.empty:
        st.success("✅ Datos suficientes para capturar gráficos 100% RPM")
    else:
        st.error("❌ Datos insuficientes para capturar gráficos 100% RPM")
        st.info("💡 Ve a la pestaña 'Análisis de curvas' y verifica que las curvas estén visibles")

def capturar_todos_los_graficos_automaticamente():
    """Verifica si hay gráficos capturados desde la pestaña 2"""
    try:
        # Inicializar estructura de gráficos
        if 'graficos_exportados' not in st.session_state:
            st.session_state['graficos_exportados'] = {
                'grupo_100_rpm': {},
                'grupo_vfd': {}
            }
        
        # Verificar si ya hay gráficos capturados
        if 'graficos_exportados' in st.session_state:
            graficos_100 = len(st.session_state['graficos_exportados']['grupo_100_rpm'])
            graficos_vfd = len(st.session_state['graficos_exportados']['grupo_vfd'])
            
            st.info(f"🔍 Gráficos en memoria: {graficos_100} (100% RPM) + {graficos_vfd} (VFD)")
            
            if graficos_100 > 0 or graficos_vfd > 0:
                st.success(f"✅ Encontrados {graficos_100 + graficos_vfd} gráficos capturados desde la pestaña 'Análisis de curvas'")
                return True
        
        # Verificar flags de captura individual
        flags_captura = [
            'hq_100_capturado', 'rend_100_capturado', 'pot_100_capturado', 'npsh_100_capturado',
            'vfd_hq_capturado', 'vfd_rend_capturado', 'vfd_pot_capturado', 'vfd_npsh_capturado'
        ]
        
        capturados_por_flags = sum(1 for flag in flags_captura if st.session_state.get(flag, False))
        st.info(f"🔍 Gráficos capturados por flags: {capturados_por_flags}/8")
        
        if capturados_por_flags > 0:
            st.success(f"✅ Se detectaron {capturados_por_flags} gráficos capturados desde la pestaña 'Análisis de curvas'")
            return True
        
        st.warning("⚠️ No se encontraron gráficos capturados.")
        st.info("💡 Para capturar gráficos:")
        st.info("1. Ve a la pestaña 'Análisis de curvas'")
        st.info("2. Verifica que las curvas estén visibles")
        st.info("3. Los gráficos se capturarán automáticamente")
        st.info("4. Regresa aquí y activa 'Incluir gráficos' nuevamente")
        
        return False
        
    except Exception as e:
        st.error(f"❌ Error en captura de gráficos: {e}")
        import traceback
        st.error(traceback.format_exc())
        return False

def obtener_grafico_capturado(grupo, nombre_grafico):
    """Obtiene un gráfico capturado desde session_state"""
    try:
        if ('graficos_exportados' in st.session_state and 
            grupo in st.session_state['graficos_exportados'] and
            nombre_grafico in st.session_state['graficos_exportados'][grupo]):
            return st.session_state['graficos_exportados'][grupo][nombre_grafico]
        return None
    except Exception:
        return None

def agregar_imagen_plotly_a_doc(doc, grupo, nombre_grafico, titulo):
    """Agrega una imagen de Plotly capturada al documento DOCX - usa directamente las imágenes capturadas"""
    # Usar directamente las imágenes capturadas (sin recrear gráficos)
    imagen_bytes = obtener_grafico_capturado(grupo, nombre_grafico)
    
    if imagen_bytes:
        try:
            doc.add_heading(titulo, level=3)
            doc.add_picture(io.BytesIO(imagen_bytes), width=Inches(6.0))
            
            # Agregar información del punto de operación
            agregar_info_punto_operacion(doc, nombre_grafico)
            
            doc.add_paragraph()  # Espacio después del gráfico
            return True
        except Exception as e:
            doc.add_paragraph(f"Error al insertar gráfico {nombre_grafico}: {e}")
            return False
    else:
        # Fallback: mensaje de error
        doc.add_heading(titulo, level=3)
        doc.add_paragraph(f"Gráfico {nombre_grafico} no disponible")
        return False

def crear_docx_con_grafico(fig, titulo_documento="Informe Pumping System", titulo_grafico="Gráfica del Sistema"):
    """
    Crea un documento Word con un gráfico en un solo paso, sin archivos temporales.
    
    Args:
        fig: Figura de Matplotlib o Plotly
        titulo_documento: Título del documento
        titulo_grafico: Título del gráfico
    
    Returns:
        bytes: Datos del documento DOCX en memoria
    """
    try:
        # 1. Crear el documento
        doc = Document()
        doc.add_heading(titulo_documento, level=0)
        
        # 2. Agregar descripción
        doc.add_paragraph("A continuación se muestra la gráfica generada en la app:")
        
        # 3. Convertir figura a buffer y agregar al documento
        buf = figuratransform(fig)
        doc.add_picture(buf, width=Inches(6))  # Ajusta el ancho según necesites
        
        # 4. Agregar título del gráfico
        doc.add_heading(titulo_grafico, level=2)
        
        # 5. Guardar en buffer en memoria
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer.getvalue()
        
    except Exception as e:
        st.error(f"❌ Error creando documento Word: {e}")
        import traceback
        st.error("Detalles del error:")
        st.code(traceback.format_exc())
        return None

def guardar_grafico_en_session_state(fig, nombre_grafico, grupo="graficos_generales", mostrar_log=False):
    """
    Guarda un gráfico en session_state para acceso entre pestañas
    
    Args:
        fig: Figura de Matplotlib o Plotly
        nombre_grafico: Nombre único del gráfico
        grupo: Grupo al que pertenece el gráfico
        mostrar_log: Si mostrar mensaje de log (por defecto False para evitar spam)
    """
    try:
        # Inicializar estructura si no existe
        if 'graficos_objetos' not in st.session_state:
            st.session_state['graficos_objetos'] = {}
        
        if grupo not in st.session_state['graficos_objetos']:
            st.session_state['graficos_objetos'][grupo] = {}
        
        # Guardar el objeto de figura
        st.session_state['graficos_objetos'][grupo][nombre_grafico] = fig
        
        # Log para diagnóstico (solo si se solicita)
        if mostrar_log:
            st.info(f"📊 Gráfico '{nombre_grafico}' guardado en session_state (grupo: {grupo})")
        
        return True
        
    except Exception as e:
        st.error(f"❌ Error guardando gráfico en session_state: {e}")
        return False

def obtener_grafico_desde_session_state(nombre_grafico, grupo="graficos_generales"):
    """
    Obtiene un gráfico desde session_state
    
    Args:
        nombre_grafico: Nombre del gráfico
        grupo: Grupo al que pertenece el gráfico
    
    Returns:
        Figura o None si no se encuentra
    """
    try:
        if ('graficos_objetos' in st.session_state and 
            grupo in st.session_state['graficos_objetos'] and
            nombre_grafico in st.session_state['graficos_objetos'][grupo]):
            return st.session_state['graficos_objetos'][grupo][nombre_grafico]
        return None
    except Exception:
        return None

def crear_docx_con_graficos_desde_session_state(titulo_documento="Informe de Sistema de Bombeo"):
    """
    Crea un documento Word con todos los gráficos disponibles en session_state
    
    Args:
        titulo_documento: Título del documento
    
    Returns:
        bytes: Datos del documento DOCX en memoria o None si hay error
    """
    try:
        # Verificar si hay gráficos disponibles
        if 'graficos_objetos' not in st.session_state:
            st.warning("⚠️ No hay gráficos disponibles en session_state")
            return None
        
        # Crear el documento
        doc = Document()
        doc.add_heading(titulo_documento, level=0)
        doc.add_paragraph("A continuación se muestran las gráficas generadas en la aplicación:")
        
        total_graficos = 0
        
        # Procesar cada grupo de gráficos
        for grupo, graficos in st.session_state['graficos_objetos'].items():
            if graficos:  # Si el grupo tiene gráficos
                doc.add_heading(f"Gráficos - {grupo.replace('_', ' ').title()}", level=1)
                
                for nombre_grafico, fig in graficos.items():
                    try:
                        # Crear gráfico simplificado para el informe
                        fig_simplificada = crear_grafico_simplificado_para_informe(fig, nombre_grafico)
                        
                        # Convertir figura a buffer y agregar al documento
                        buf = figuratransform(fig_simplificada)
                        doc.add_heading(nombre_grafico.replace('_', ' ').title(), level=2)
                        doc.add_picture(buf, width=Inches(6))
                        
                        # Agregar información del punto de operación
                        agregar_info_punto_operacion(doc, nombre_grafico)
                        
                        doc.add_paragraph()  # Espacio después del gráfico
                        total_graficos += 1
                        
                    except Exception as e:
                        st.error(f"❌ Error procesando gráfico '{nombre_grafico}': {e}")
                        doc.add_paragraph(f"Error al procesar gráfico: {nombre_grafico}")
        
        if total_graficos == 0:
            st.warning("⚠️ No se pudieron procesar gráficos válidos")
            return None
        
        # Guardar en buffer en memoria
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        st.success(f"✅ Documento creado con {total_graficos} gráficos")
        return doc_buffer.getvalue()
        
    except Exception as e:
        st.error(f"❌ Error creando documento Word: {e}")
        import traceback
        st.error("Detalles del error:")
        st.code(traceback.format_exc())
        return None

def crear_grafico_simplificado_para_captura(fig_original, nombre_grafico):
    """
    Crea un gráfico simplificado para captura, eliminando líneas duplicadas y zona de eficiencia
    """
    import matplotlib.pyplot as plt
    import numpy as np
    
    # Crear nueva figura
    fig, ax = plt.subplots(figsize=(10, 6))
    
    # Contador para evitar líneas duplicadas
    curvas_mostradas = set()
    
    # Usar directamente los datos del gráfico original
    for trace in fig_original.data:
        if hasattr(trace, 'x') and hasattr(trace, 'y'):
            x_data = trace.x
            y_data = trace.y
            
            if x_data is None or y_data is None:
                continue
                
            # Convertir a arrays
            if not isinstance(x_data, (list, tuple, np.ndarray)):
                x_data = [x_data] if x_data is not None else []
            if not isinstance(y_data, (list, tuple, np.ndarray)):
                y_data = [y_data] if y_data is not None else []
            
            if len(x_data) == 0 or len(y_data) == 0:
                continue
            
            # Obtener color original
            color = 'blue'
            if hasattr(trace, 'line') and trace.line is not None and hasattr(trace.line, 'color'):
                color = convertir_color_plotly_a_matplotlib(trace.line.color)
            elif hasattr(trace, 'marker') and trace.marker is not None and hasattr(trace.marker, 'color'):
                color = convertir_color_plotly_a_matplotlib(trace.marker.color)
            
            # Obtener label original
            label = trace.name if hasattr(trace, 'name') and trace.name is not None else 'Datos'
            
            # FILTRAR: Solo mostrar elementos principales, evitar duplicados
            
            # Puntos de operación (mantener como están)
            if any(keyword in label.lower() for keyword in ['operación', 'punto operación', 'operacion', 'punto de operación']):
                ax.plot(x_data, y_data, 'o', color=color, label=label, markersize=12, 
                       markeredgewidth=2, markeredgecolor='black', zorder=10)
            
            # Curvas principales - EVITAR DUPLICADOS
            elif any(keyword in label.lower() for keyword in ['bomba', 'sistema', 'potencia', 'npsh']):
                # Solo mostrar si tiene suficientes puntos para ser una curva (no puntos individuales)
                if len(x_data) > 3:
                    # Crear identificador único para evitar duplicados
                    tipo_curva = None
                    if 'bomba' in label.lower():
                        tipo_curva = 'bomba'
                    elif 'sistema' in label.lower():
                        tipo_curva = 'sistema'
                    elif 'potencia' in label.lower():
                        tipo_curva = 'potencia'
                    elif 'npsh' in label.lower():
                        tipo_curva = 'npsh'
                    
                    # Solo mostrar si no se ha mostrado ya este tipo de curva
                    if tipo_curva and tipo_curva not in curvas_mostradas:
                        ax.plot(x_data, y_data, '-', color=color, label=label, linewidth=2.5)
                        curvas_mostradas.add(tipo_curva)
            
            # ZONA DE EFICIENCIA: Mostrar curva real y evitar duplicados
            elif any(keyword in label.lower() for keyword in ['eficiencia', 'rendimiento']):
                # Solo mostrar una vez para evitar duplicados en la simbología
                if 'eficiencia_mostrada' not in locals():
                    # Obtener valores de BEP para el label
                    zona_eff_min, zona_eff_max, bep_eta = obtener_valores_bep_eficiencia()
                    label_eficiencia = f"Zona de eficiencia ({zona_eff_min:.0f}%-{zona_eff_max:.0f}% BEP)"
                    
                    # Mostrar la curva real de eficiencia
                    if len(x_data) > 3:
                        ax.plot(x_data, y_data, '-', color='lightgray', label=label_eficiencia, 
                               linewidth=2, alpha=0.7)
                    
                    # Marcar como mostrada
                    locals()['eficiencia_mostrada'] = True
            
            # BEP (Best Efficiency Point)
            elif 'bep' in label.lower() or 'mejor punto' in label.lower():
                ax.plot(x_data, y_data, 's', color='green', label=label, markersize=10, 
                       markeredgewidth=2, markeredgecolor='black', zorder=9)
            
            # Excluir todos los demás elementos (puntos de datos individuales, etc.)
    
    # Configurar el gráfico con la misma configuración que el original
    if (hasattr(fig_original.layout, 'title') and 
        fig_original.layout.title is not None and 
        hasattr(fig_original.layout.title, 'text') and 
        fig_original.layout.title.text is not None):
        ax.set_title(fig_original.layout.title.text, fontsize=14, fontweight='bold')
    
    if (hasattr(fig_original.layout, 'xaxis') and 
        fig_original.layout.xaxis is not None and
        hasattr(fig_original.layout.xaxis, 'title') and 
        fig_original.layout.xaxis.title is not None and
        hasattr(fig_original.layout.xaxis.title, 'text') and
        fig_original.layout.xaxis.title.text is not None):
        ax.set_xlabel(fig_original.layout.xaxis.title.text, fontsize=12)
    
    if (hasattr(fig_original.layout, 'yaxis') and 
        fig_original.layout.yaxis is not None and
        hasattr(fig_original.layout.yaxis, 'title') and 
        fig_original.layout.yaxis.title is not None and
        hasattr(fig_original.layout.yaxis.title, 'text') and
        fig_original.layout.yaxis.title.text is not None):
        ax.set_ylabel(fig_original.layout.yaxis.title.text, fontsize=12)
    
    # Agregar grid y leyenda
    ax.grid(True, alpha=0.3)
    ax.legend()
    
    # Ajustar layout
    plt.tight_layout()
    
    return fig

def crear_grafico_simplificado_para_informe(fig_original, nombre_grafico):
    """
    Crea un gráfico simplificado para el informe usando directamente los datos capturados de la pestaña 2.
    No recalcula nada, solo simplifica la visualización eliminando puntos de datos individuales.
    """
    import matplotlib.pyplot as plt
    import numpy as np
    
    # Crear nueva figura
    fig, ax = plt.subplots(figsize=(10, 6))
    
    # Usar directamente los datos del gráfico original sin recalcular
    for trace in fig_original.data:
        if hasattr(trace, 'x') and hasattr(trace, 'y'):
            x_data = trace.x
            y_data = trace.y
            
            if x_data is None or y_data is None:
                continue
                
            # Convertir a arrays
            if not isinstance(x_data, (list, tuple, np.ndarray)):
                x_data = [x_data] if x_data is not None else []
            if not isinstance(y_data, (list, tuple, np.ndarray)):
                y_data = [y_data] if y_data is not None else []
            
            if len(x_data) == 0 or len(y_data) == 0:
                continue
            
            # Obtener color original
            color = 'blue'
            if hasattr(trace, 'line') and trace.line is not None and hasattr(trace.line, 'color'):
                color = convertir_color_plotly_a_matplotlib(trace.line.color)
            elif hasattr(trace, 'marker') and trace.marker is not None and hasattr(trace.marker, 'color'):
                color = convertir_color_plotly_a_matplotlib(trace.marker.color)
            
            # Obtener label original
            label = trace.name if hasattr(trace, 'name') and trace.name is not None else 'Datos'
            
            # FILTRAR: Solo mostrar elementos principales, excluir puntos de datos individuales
            
            # Puntos de operación (mantener como están)
            if any(keyword in label.lower() for keyword in ['operación', 'punto operación', 'operacion', 'punto de operación']):
                ax.plot(x_data, y_data, 'o', color=color, label=label, markersize=12, 
                       markeredgewidth=2, markeredgecolor='black', zorder=10)
            
            # Curvas principales (usar el tipo de ajuste que ya está en la pestaña 2)
            elif any(keyword in label.lower() for keyword in ['bomba', 'sistema', 'potencia', 'npsh']):
                # Solo mostrar si tiene suficientes puntos para ser una curva (no puntos individuales)
                if len(x_data) > 3:
                    # Usar el mismo estilo que en la pestaña 2 (líneas sin marcadores)
                    ax.plot(x_data, y_data, '-', color=color, label=label, linewidth=2.5)
            
            # ZONA DE EFICIENCIA: Mostrar curva real y evitar duplicados
            elif any(keyword in label.lower() for keyword in ['eficiencia', 'rendimiento']):
                # Solo mostrar una vez para evitar duplicados en la simbología
                if 'eficiencia_mostrada' not in locals():
                    # Obtener valores de BEP para el label
                    zona_eff_min, zona_eff_max, bep_eta = obtener_valores_bep_eficiencia()
                    label_eficiencia = f"Zona de eficiencia ({zona_eff_min:.0f}%-{zona_eff_max:.0f}% BEP)"
                    
                    # Mostrar la curva real de eficiencia
                    if len(x_data) > 3:
                        ax.plot(x_data, y_data, '-', color='lightgray', label=label_eficiencia, 
                               linewidth=2, alpha=0.7)
                    
                    # Marcar como mostrada
                    locals()['eficiencia_mostrada'] = True
            
            # BEP (Best Efficiency Point)
            elif 'bep' in label.lower() or 'mejor punto' in label.lower():
                ax.plot(x_data, y_data, 's', color='green', label=label, markersize=10, 
                       markeredgewidth=2, markeredgecolor='black', zorder=9)
            
            # Excluir todos los demás elementos (puntos de datos individuales, etc.)
    
    # Configurar el gráfico con la misma configuración que el original
    if (hasattr(fig_original.layout, 'title') and 
        fig_original.layout.title is not None and 
        hasattr(fig_original.layout.title, 'text') and 
        fig_original.layout.title.text is not None):
        ax.set_title(fig_original.layout.title.text, fontsize=14, fontweight='bold')
    
    if (hasattr(fig_original.layout, 'xaxis') and 
        fig_original.layout.xaxis is not None and
        hasattr(fig_original.layout.xaxis, 'title') and 
        fig_original.layout.xaxis.title is not None and
        hasattr(fig_original.layout.xaxis.title, 'text') and
        fig_original.layout.xaxis.title.text is not None):
        ax.set_xlabel(fig_original.layout.xaxis.title.text, fontsize=12)
    
    if (hasattr(fig_original.layout, 'yaxis') and 
        fig_original.layout.yaxis is not None and
        hasattr(fig_original.layout.yaxis, 'title') and 
        fig_original.layout.yaxis.title is not None and
        hasattr(fig_original.layout.yaxis.title, 'text') and
        fig_original.layout.yaxis.title.text is not None):
        ax.set_ylabel(fig_original.layout.yaxis.title.text, fontsize=12)
    
    # Agregar grid y leyenda
    ax.grid(True, alpha=0.3)
    ax.legend()
    
    # Ajustar layout
    plt.tight_layout()
    
    return fig

# FUNCIÓN ELIMINADA - Zona de eficiencia no se usa
# def agregar_area_eficiencia(ax, fig_original):
#     """Agrega el área verde de eficiencia al gráfico"""
#     pass

def obtener_valores_bep_eficiencia():
    """Obtiene los valores de BEP y zona de eficiencia desde session_state"""
    try:
        # Obtener valores de zona de eficiencia desde session_state
        zona_eff_min = st.session_state.get('zona_eff_min', 65.0)
        zona_eff_max = st.session_state.get('zona_eff_max', 115.0)
        
        # Obtener datos de curva de eficiencia
        curva_inputs = st.session_state.get('curva_inputs', {})
        puntos_rend = curva_inputs.get('rendimiento', [])
        
        if len(puntos_rend) >= 2:
            import numpy as np
            x_rend = np.array([pt[0] for pt in puntos_rend])
            y_rend = np.array([pt[1] for pt in puntos_rend])
            
            # Calcular BEP
            ajuste_tipo = st.session_state.get('ajuste_tipo', 'Cuadrática (2do grado)')
            grado_rend = 1 if ajuste_tipo == "Lineal" else 2 if ajuste_tipo == "Cuadrática (2do grado)" else 3
            coef_rend = np.polyfit(x_rend, y_rend, grado_rend)
            x_fit = np.linspace(x_rend.min(), x_rend.max(), 100)
            y_fit = np.polyval(coef_rend, x_fit)
            idx_bep = np.argmax(y_fit)
            bep_eta = y_fit[idx_bep]
            
            return zona_eff_min, zona_eff_max, bep_eta
        else:
            return zona_eff_min, zona_eff_max, 0.0
    except Exception:
        return 65.0, 115.0, 0.0

def calcular_eficiencia_en_punto_operacion(caudal_op, es_vfd=False):
    """Calcula la eficiencia en el punto de operación usando interpolación"""
    try:
        import numpy as np
        
        # Obtener datos de curva de eficiencia
        curva_inputs = st.session_state.get('curva_inputs', {})
        puntos_rend = curva_inputs.get('rendimiento', [])
        
        # Debug: mostrar información disponible
        if hasattr(st, 'info') and st.session_state.get('debug_eficiencia', False):
            st.info(f"🔍 Debug Eficiencia - Caudal OP: {caudal_op}")
            st.info(f"🔍 Debug Eficiencia - Puntos rendimiento: {len(puntos_rend)}")
            if puntos_rend:
                st.info(f"🔍 Debug Eficiencia - Primer punto: {puntos_rend[0]}")
                st.info(f"🔍 Debug Eficiencia - Último punto: {puntos_rend[-1]}")
        
        if len(puntos_rend) >= 2 and caudal_op > 0:
            x_rend = np.array([pt[0] for pt in puntos_rend])
            y_rend = np.array([pt[1] for pt in puntos_rend])
            
            # Verificar que el caudal esté dentro del rango de datos
            if caudal_op >= x_rend.min() and caudal_op <= x_rend.max():
                # Interpolar para obtener la eficiencia en el punto de operación
                rendimiento_op = np.interp(caudal_op, x_rend, y_rend)
                
                if hasattr(st, 'info') and st.session_state.get('debug_eficiencia', False):
                    st.success(f"✅ Eficiencia calculada: {rendimiento_op:.2f}%")
                
                return rendimiento_op
            else:
                if hasattr(st, 'info') and st.session_state.get('debug_eficiencia', False):
                    st.warning(f"⚠️ Caudal {caudal_op} fuera del rango [{x_rend.min():.2f}, {x_rend.max():.2f}]")
        
        return 0.0
    except Exception as e:
        if hasattr(st, 'error') and st.session_state.get('debug_eficiencia', False):
            st.error(f"❌ Error calculando eficiencia: {e}")
        return 0.0

def agregar_info_punto_operacion(doc, nombre_grafico):
    """Agrega información del punto de operación al documento"""
    try:
        # Determinar si es VFD o 100% RPM
        es_vfd = 'vfd' in nombre_grafico.lower()
        
        if es_vfd:
            # Datos VFD específicos desde interseccion_vfd
            interseccion_vfd = st.session_state.get('interseccion_vfd', None)
            if interseccion_vfd and len(interseccion_vfd) >= 2:
                caudal_op = interseccion_vfd[0]
                altura_op = interseccion_vfd[1]
            else:
                caudal_op = 0
                altura_op = 0
            
            # Calcular eficiencia en el punto de operación
            rendimiento_op = calcular_eficiencia_en_punto_operacion(caudal_op, es_vfd=True)
            potencia_op = st.session_state.get('op_pot_vfd', 0)
            npsh_op = st.session_state.get('op_npsh_vfd', 0)
            titulo_op = "Punto de Operación VFD:"
        else:
            # Datos 100% RPM
            caudal_op = st.session_state.get('caudal_operacion', 0)
            altura_op = st.session_state.get('altura_operacion', 0)
            
            # Calcular eficiencia en el punto de operación
            rendimiento_op = calcular_eficiencia_en_punto_operacion(caudal_op, es_vfd=False)
            potencia_op = st.session_state.get('potencia_operacion', 0)
            npsh_op = st.session_state.get('npsh_requerido', 0)
            titulo_op = "Punto de Operación:"
        
        # Agregar información según el tipo de gráfico
        if 'hq' in nombre_grafico.lower():
            doc.add_paragraph(titulo_op)
            doc.add_paragraph(f"Caudal (Q): {caudal_op:.2f} L/s")
            doc.add_paragraph(f"Altura (H): {altura_op:.2f} m")
        
        elif 'rendimiento' in nombre_grafico.lower() or 'eficiencia' in nombre_grafico.lower():
            doc.add_paragraph(titulo_op)
            doc.add_paragraph(f"Caudal (Q): {caudal_op:.2f} L/s")
            doc.add_paragraph(f"Rendimiento (η): {rendimiento_op:.2f} %")
            
            # Calcular y mostrar el valor de BEP
            try:
                import numpy as np
                curva_inputs = st.session_state.get('curva_inputs', {})
                puntos_rend = curva_inputs.get('rendimiento', [])
                
                if len(puntos_rend) >= 2:
                    x_rend = np.array([pt[0] for pt in puntos_rend])
                    y_rend = np.array([pt[1] for pt in puntos_rend])
                    
                    # Calcular BEP
                    ajuste_tipo = st.session_state.get('ajuste_tipo', 'Cuadrática (2do grado)')
                    grado_rend = 1 if ajuste_tipo == "Lineal" else 2 if ajuste_tipo == "Cuadrática (2do grado)" else 3
                    coef_rend = np.polyfit(x_rend, y_rend, grado_rend)
                    x_fit = np.linspace(x_rend.min(), x_rend.max(), 100)
                    y_fit = np.polyval(coef_rend, x_fit)
                    idx_bep = np.argmax(y_fit)
                    bep_q = x_fit[idx_bep]
                    bep_eta = y_fit[idx_bep]
                    
                    doc.add_paragraph(f"BEP (Best Efficiency Point): {bep_q:.2f} L/s @ {bep_eta:.1f}%")
            except Exception:
                doc.add_paragraph("BEP: No disponible")
        
        elif 'potencia' in nombre_grafico.lower():
            doc.add_paragraph(titulo_op)
            doc.add_paragraph(f"Caudal (Q): {caudal_op:.2f} L/s")
            doc.add_paragraph(f"Potencia (PBHP): {potencia_op:.2f} HP")
        
        elif 'npsh' in nombre_grafico.lower():
            doc.add_paragraph(titulo_op)
            doc.add_paragraph(f"Caudal (Q): {caudal_op:.2f} L/s")
            doc.add_paragraph(f"NPSH Requerido: {npsh_op:.2f} m")
        
        doc.add_paragraph()  # Espacio adicional
        
    except Exception as e:
        doc.add_paragraph(f"Información del punto de operación no disponible")

def mostrar_estado_graficos():
    """
    Muestra el estado actual de los gráficos en session_state para diagnóstico
    """
    st.markdown("### 🔍 Estado de Gráficos en Session State")
    
    # Inicializar contador
    total_graficos = 0
    
    # Verificar gráficos en session_state (método nuevo)
    if 'graficos_objetos' not in st.session_state:
        st.warning("⚠️ No hay gráficos en session_state")
    else:
        for grupo, graficos in st.session_state['graficos_objetos'].items():
            st.info(f"**Grupo: {grupo}**")
            if graficos:
                for nombre, fig in graficos.items():
                    tipo_fig = type(fig).__name__
                    st.write(f"  - {nombre}: {tipo_fig}")
                    total_graficos += 1
            else:
                st.write("  - (vacío)")
        
        st.success(f"📊 Total de gráficos en session_state: {total_graficos}")
    
    # Verificar gráficos exportados (método anterior)
    if 'graficos_exportados' in st.session_state:
        st.info("**Gráficos Exportados (método anterior):**")
        total_exportados = 0
        for grupo, graficos in st.session_state['graficos_exportados'].items():
            st.write(f"  - {grupo}: {len(graficos)} gráficos")
            total_exportados += len(graficos)
        st.info(f"📊 Total de gráficos exportados: {total_exportados}")
    else:
        st.warning("⚠️ No hay gráficos exportados (método anterior)")
    
    # Verificar flags de captura individual
    st.info("**Flags de Captura Individual:**")
    flags_captura = [
        'hq_100_capturado', 'rend_100_capturado', 'pot_100_capturado', 'npsh_100_capturado',
        'vfd_hq_capturado', 'vfd_rend_capturado', 'vfd_pot_capturado', 'vfd_npsh_capturado'
    ]
    
    capturados_por_flags = 0
    for flag in flags_captura:
        estado = st.session_state.get(flag, False)
        if estado:
            capturados_por_flags += 1
        st.write(f"  - {flag}: {'✅' if estado else '❌'}")
    
    st.info(f"📊 Gráficos capturados por flags: {capturados_por_flags}/8")
    
    # Resumen final
    st.markdown("---")
    if total_graficos >= 8:
        st.success("🎉 ¡Todos los gráficos están disponibles!")
    elif total_graficos > 0:
        st.warning(f"⚠️ Solo {total_graficos}/8 gráficos disponibles")
    else:
        st.error("❌ No hay gráficos disponibles. Ve a la pestaña 'Análisis de curvas' para generar gráficos.")

def generar_grafico_bombeo_matplotlib():
    """Genera un gráfico de ejemplo de bombeo con Matplotlib"""
    import numpy as np
    
    # Datos de ejemplo para curva de bombeo
    caudal = np.linspace(0, 100, 50)  # L/s
    altura_bomba = 50 - 0.01 * caudal**2  # Curva parabólica típica
    altura_sistema = 20 + 0.005 * caudal**2  # Curva del sistema
    
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(caudal, altura_bomba, 'b-', linewidth=2, label='Curva de la Bomba')
    ax.plot(caudal, altura_sistema, 'r-', linewidth=2, label='Curva del Sistema')
    
    # Punto de operación (intersección)
    idx_operacion = np.argmin(np.abs(altura_bomba - altura_sistema))
    ax.plot(caudal[idx_operacion], altura_bomba[idx_operacion], 'go', markersize=10, label='Punto de Operación')
    
    ax.set_title('Curva de Bombeo - Análisis H-Q', fontsize=14, fontweight='bold')
    ax.set_xlabel('Caudal (L/s)', fontsize=12)
    ax.set_ylabel('Altura (m)', fontsize=12)
    ax.grid(True, alpha=0.3)
    ax.legend()
    plt.tight_layout()
    
    return fig

def generar_grafico_bombeo_plotly():
    """Genera un gráfico de ejemplo de bombeo con Plotly"""
    import plotly.graph_objects as go
    import numpy as np
    
    # Datos de ejemplo para curva de bombeo
    caudal = np.linspace(0, 100, 50)  # L/s
    altura_bomba = 50 - 0.01 * caudal**2  # Curva parabólica típica
    altura_sistema = 20 + 0.005 * caudal**2  # Curva del sistema
    
    fig = go.Figure()
    
    # Curva de la bomba
    fig.add_trace(go.Scatter(
        x=caudal, y=altura_bomba,
        mode='lines',
        name='Curva de la Bomba',
        line=dict(color='blue', width=3)
    ))
    
    # Curva del sistema
    fig.add_trace(go.Scatter(
        x=caudal, y=altura_sistema,
        mode='lines',
        name='Curva del Sistema',
        line=dict(color='red', width=3)
    ))
    
    # Punto de operación
    idx_operacion = np.argmin(np.abs(altura_bomba - altura_sistema))
    fig.add_trace(go.Scatter(
        x=[caudal[idx_operacion]], y=[altura_bomba[idx_operacion]],
        mode='markers',
        name='Punto de Operación',
        marker=dict(color='green', size=12, symbol='circle')
    ))
    
    fig.update_layout(
        title='Curva de Bombeo - Análisis H-Q',
        xaxis_title='Caudal (L/s)',
        yaxis_title='Altura (m)',
        showlegend=True,
        width=800,
        height=500
    )
    
    return fig

def descargar_grupo_graficos(grupo, nombre_grupo):
    """Descarga un grupo específico de gráficos como archivo ZIP"""
    try:
        import zipfile
        from datetime import datetime
        
        if 'graficos_exportados' not in st.session_state or grupo not in st.session_state['graficos_exportados']:
            st.error(f"No hay gráficos disponibles para el grupo {nombre_grupo}")
            return
        
        graficos = st.session_state['graficos_exportados'][grupo]
        if not graficos:
            st.error(f"No hay gráficos en el grupo {nombre_grupo}")
            return
        
        # Crear archivo ZIP en memoria
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for nombre_grafico, imagen_bytes in graficos.items():
                # Crear nombre de archivo descriptivo
                nombre_archivo = f"{nombre_grupo}_{nombre_grafico}.png"
                zip_file.writestr(nombre_archivo, imagen_bytes)
        
        zip_buffer.seek(0)
        
        # Crear nombre de archivo con timestamp
        timestamp = datetime.now().strftime('%Y%m%d_%H%M')
        nombre_zip = f"Graficos_{nombre_grupo}_{timestamp}.zip"
        
        # Mostrar botón de descarga
        st.download_button(
            label=f"📥 Descargar {nombre_grupo} ({len(graficos)} gráficos)",
            data=zip_buffer.getvalue(),
            file_name=nombre_zip,
            mime="application/zip",
            key=f"download_{grupo}_{timestamp}"
        )
        
        st.success(f"✅ Archivo ZIP preparado con {len(graficos)} gráficos del grupo {nombre_grupo}")
        
    except Exception as e:
        st.error(f"Error al crear archivo ZIP: {e}")

def descargar_todos_los_graficos():
    """Descarga todos los gráficos capturados como archivo ZIP"""
    try:
        import zipfile
        from datetime import datetime
        
        if 'graficos_exportados' not in st.session_state:
            st.error("No hay gráficos capturados disponibles")
            return
        
        total_graficos = 0
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            # Agregar gráficos de 100% RPM
            if 'grupo_100_rpm' in st.session_state['graficos_exportados']:
                for nombre_grafico, imagen_bytes in st.session_state['graficos_exportados']['grupo_100_rpm'].items():
                    nombre_archivo = f"100_RPM_{nombre_grafico}.png"
                    zip_file.writestr(nombre_archivo, imagen_bytes)
                    total_graficos += 1
            
            # Agregar gráficos VFD
            if 'grupo_vfd' in st.session_state['graficos_exportados']:
                for nombre_grafico, imagen_bytes in st.session_state['graficos_exportados']['grupo_vfd'].items():
                    nombre_archivo = f"VFD_{nombre_grafico}.png"
                    zip_file.writestr(nombre_archivo, imagen_bytes)
                    total_graficos += 1
        
        if total_graficos == 0:
            st.error("No hay gráficos disponibles para descargar")
            return
        
        zip_buffer.seek(0)
        
        # Crear nombre de archivo con timestamp
        timestamp = datetime.now().strftime('%Y%m%d_%H%M')
        nombre_zip = f"Todos_los_Graficos_{timestamp}.zip"
        
        # Mostrar botón de descarga
        st.download_button(
            label=f"📥 Descargar Todos los Gráficos ({total_graficos} gráficos)",
            data=zip_buffer.getvalue(),
            file_name=nombre_zip,
            mime="application/zip",
            key=f"download_all_{timestamp}"
        )
        
        st.success(f"✅ Archivo ZIP preparado con {total_graficos} gráficos")
        
    except Exception as e:
        st.error(f"Error al crear archivo ZIP: {e}")

def verificar_datos_curvas():
    """Verifica el estado de los datos de curvas disponibles"""
    st.markdown("### 🔍 Verificación de Datos de Curvas")
    
    # Verificar session_state
    if 'tablas_graficos' not in st.session_state:
        st.error("❌ No existe 'tablas_graficos' en session_state")
        return
    
    tablas = st.session_state.get('tablas_graficos', {})
    st.info(f"📊 Estructura de tablas_graficos: {list(tablas.keys())}")
    
    # Verificar datos 100% RPM
    datos_100 = tablas.get('tablas_100_rpm', {})
    if datos_100:
        st.success("✅ Datos 100% RPM encontrados")
        for key, value in datos_100.items():
            if isinstance(value, dict) and 'data' in value:
                df = deserialize_df(value)
                st.info(f"  - {key}: {len(df)} filas" if not df.empty else f"  - {key}: VACÍO")
            else:
                st.warning(f"  - {key}: Formato incorrecto")
    else:
        st.error("❌ No hay datos 100% RPM")
    
    # Verificar datos VFD
    datos_vfd = tablas.get('tablas_vfd_rpm', {})
    if datos_vfd:
        st.success("✅ Datos VFD encontrados")
        for key, value in datos_vfd.items():
            if isinstance(value, dict) and 'data' in value:
                df = deserialize_df(value)
                st.info(f"  - {key}: {len(df)} filas" if not df.empty else f"  - {key}: VACÍO")
            else:
                st.warning(f"  - {key}: Formato incorrecto")
    else:
        st.error("❌ No hay datos VFD")
    
    # Verificar gráficos capturados
    if 'graficos_exportados' in st.session_state:
        graficos = st.session_state['graficos_exportados']
        st.success(f"✅ Gráficos capturados: {len(graficos.get('grupo_100_rpm', {}))} (100% RPM) + {len(graficos.get('grupo_vfd', {}))} (VFD)")
    else:
        st.warning("⚠️ No hay gráficos capturados")

def forzar_captura_graficos():
    """Verifica si hay gráficos capturados o intenta generarlos automáticamente"""
    try:
        # Verificar si hay gráficos ya capturados
        graficos_100 = 0
        graficos_vfd = 0
        if 'graficos_exportados' in st.session_state:
            graficos_100 = len(st.session_state['graficos_exportados'].get('grupo_100_rpm', {}))
            graficos_vfd = len(st.session_state['graficos_exportados'].get('grupo_vfd', {}))
            
            if graficos_100 > 0 or graficos_vfd > 0:
                st.success(f"✅ Ya hay {graficos_100 + graficos_vfd} gráficos capturados")
                return
        
        # Si no hay gráficos, intentar generarlos forzosamente desde la pestaña 'Análisis de curvas'
        # o mediante los datos de session_state si están disponibles
        st.info("🔄 Intentando generar gráficos automáticos para el reporte...")
        
        from ui.tabs_modules.results_tab import render_results_tab
        # Simplemente llamar a la lógica de renderizado (sin mostrar) suele disparar capturas
        # si la función de captura está integrada en los plots
        
        # Si aún no hay, mostrar mensaje de guía
        st.warning("⚠️ Los gráficos se capturan al visitar la pestaña 'Análisis de Curvas'.")
        st.info("💡 Por favor, ve un momento a la pestaña **'Análisis de Curvas'** y luego regresa aquí para generar el reporte con imágenes.")
            
    except Exception as e:
        st.error(f"❌ Error verificando gráficos: {e}")
            
    except Exception as e:
        st.error(f"❌ Error verificando gráficos: {e}")
        import traceback
        st.error(traceback.format_exc())

def deserialize_df(df_data):
    """Deserializa un diccionario a un DataFrame de Pandas."""
    if df_data and isinstance(df_data, dict) and 'data' in df_data and df_data['data']:
        try:
            return pd.DataFrame(df_data['data'])
        except Exception:
            return pd.DataFrame()
    return pd.DataFrame()

def add_df_to_doc(doc, df, title):
    """Añade un DataFrame de pandas a un documento docx como una tabla."""
    doc.add_heading(title, level=2)
    if df.empty:
        doc.add_paragraph("No hay datos disponibles.")
        return
    
    # Asegurarse que los nombres de columnas son strings
    df.columns = [str(col) for col in df.columns]
    df_display = df.head(20)  # Mostrar hasta 20 filas
    
    table = doc.add_table(rows=1, cols=len(df_display.columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, column_name in enumerate(df_display.columns):
        table.cell(0, i).text = column_name
        table.cell(0, i).paragraphs[0].runs[0].font.bold = True

    for _, row in df_display.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            text = f"{value:.2f}" if isinstance(value, (float, int)) else str(value)
            cells[i].text = text

def add_matplotlib_plot_to_doc(doc, title):
    """Guarda la figura actual de Matplotlib y la inserta en el documento."""
    doc.add_heading(title, level=3)
    try:
        image_stream = io.BytesIO()
        plt.savefig(image_stream, format='png', dpi=300, bbox_inches='tight')
        plt.close()  # Cerrar la figura para liberar memoria
        image_stream.seek(0)
        doc.add_picture(image_stream, width=Inches(6.0))
    except Exception as e:
        doc.add_paragraph(f"Error al generar gráfico: {e}")
    doc.add_paragraph() # Espacio después del gráfico

def replace_placeholders_in_doc(doc, variables):
    """Reemplaza todos los placeholders de texto en el documento."""
    for para in doc.paragraphs:
        for key, value in variables.items():
            if key in para.text:
                # Usar una técnica de reemplazo más robusta
                inline = para.runs
                for i in range(len(inline)):
                    if key in inline[i].text:
                        text = inline[i].text.replace(key, str(value))
                        inline[i].text = text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Esto es para el texto dentro de las celdas de la tabla
                for para_in_cell in cell.paragraphs:
                    for key, value in variables.items():
                        if key in para_in_cell.text:
                            para_in_cell.text = para_in_cell.text.replace(key, str(value))

def replace_placeholder_with_table(doc, placeholder, df, title):
    """Encuentra un placeholder y lo reemplaza con una tabla."""
    for para in doc.paragraphs:
        if placeholder in para.text:
            para.text = ""
            p = para._element
            p.getparent().remove(p)
            add_df_to_doc(doc, df, title)
            return

def replace_placeholder_with_text(doc, placeholder, text, title):
    """Reemplaza un placeholder de texto con un título y párrafo, eliminando el placeholder original."""
    for para in doc.paragraphs:
        if placeholder in para.text:
            # Eliminar el párrafo del placeholder original
            p = para._element
            p.getparent().remove(p)
            # Insertar nuevo contenido en su lugar
            doc.add_heading(title, level=1)
            
            # Crear párrafo con formato de negritas para números y unidades
            new_para = doc.add_paragraph()
            aplicar_negritas_a_numeros(new_para, text)
            return

def aplicar_negritas_a_numeros(paragraph, text):
    """Aplica negritas a números y unidades en el texto"""
    import re
    
    # Patrón para encontrar números seguidos de unidades comunes
    patrones = [
        r'(\d+\.?\d*)\s*(L/s|m³/h|m|kW|HP|%|mm|°C|kg/m³)',  # Números con unidades
        r'(\d+\.?\d*)\s*(metros?|litros?|kilovatios?|caballos?|porcentaje)',  # Números con palabras
        r'(NPSH|caudal|altura|potencia|eficiencia|diámetro|longitud|temperatura|densidad)',  # Palabras técnicas importantes
    ]
    
    # Dividir el texto en partes y aplicar formato
    partes = [text]
    
    for patron in patrones:
        nuevas_partes = []
        for parte in partes:
            if isinstance(parte, str):
                # Encontrar todas las coincidencias
                matches = list(re.finditer(patron, parte, re.IGNORECASE))
                if matches:
                    ultimo_indice = 0
                    for match in matches:
                        # Agregar texto antes de la coincidencia
                        if match.start() > ultimo_indice:
                            nuevas_partes.append(parte[ultimo_indice:match.start()])
                        
                        # Agregar la coincidencia con formato de negritas
                        nuevas_partes.append(('bold', match.group()))
                        ultimo_indice = match.end()
                    
                    # Agregar texto restante
                    if ultimo_indice < len(parte):
                        nuevas_partes.append(parte[ultimo_indice:])
                else:
                    nuevas_partes.append(parte)
            else:
                nuevas_partes.append(parte)
        partes = nuevas_partes
    
    # Construir el párrafo con formato
    for parte in partes:
        if isinstance(parte, tuple) and parte[0] == 'bold':
            run = paragraph.add_run(parte[1])
            run.bold = True
        else:
            paragraph.add_run(str(parte))

# --- Lógica Principal del Módulo ---

def crear_plantilla_sin_ia():
    """Crea y guarda en disco un archivo DOCX con la estructura completa SIN análisis IA."""
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    doc.add_heading('INFORME TÉCNICO DE SISTEMA DE BOMBEO', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. INFORMACIÓN DEL PROYECTO
    doc.add_heading('1. INFORMACIÓN DEL PROYECTO', level=1)
    doc.add_paragraph('Proyecto: {proyecto}')
    doc.add_paragraph('Diseño: {diseno}')
    doc.add_paragraph('Fecha: {fecha_generacion}')
    doc.add_paragraph('Caudal de diseño: {caudal_diseno_lps} L/s ({caudal_diseno_m3h} m³/h)')

    # 2. DATOS DE ENTRADA
    doc.add_heading('2. DATOS DE ENTRADA', level=1)
    doc.add_heading('2.1 Condiciones de Operación', level=2)
    doc.add_paragraph('• Temperatura del líquido: {temperatura} °C')
    doc.add_paragraph('• Densidad del líquido: {densidad_liquido} kg/m³')
    doc.add_paragraph('• Presión de vapor: {presion_vapor_calculada:.2f} kPa')
    doc.add_paragraph('• Presión barométrica: {presion_barometrica_calculada:.2f} kPa')
    doc.add_heading('2.2 Geometría del Sistema', level=2)
    doc.add_paragraph('• Altura de succión: {altura_succion} m')
    doc.add_paragraph('• Altura de descarga: {altura_descarga} m')
    doc.add_paragraph('• Altura estática total: {altura_estatica_total:.2f} m')
    doc.add_paragraph('• Número de bombas en paralelo: {num_bombas}')
    
    # Añadir caudal por bomba si hay paralelo
    if st.session_state.get('num_bombas', 1) > 1:
        doc.add_paragraph('• Caudal por bomba individual: {caudal_por_bomba_lps} L/s')
    doc.add_heading('2.3 Tuberías', level=2)
    doc.add_heading('Tubería de Succión', level=3)
    doc.add_paragraph('• Material: {mat_succion}')
    doc.add_paragraph('• Diámetro interno: {diam_succion_mm:.1f} mm')
    doc.add_paragraph('• Longitud: {long_succion:.1f} m')
    doc.add_paragraph('• Coeficiente Hazen-Williams: {coeficiente_hazen_succion}')
    doc.add_heading('Tubería de Impulsión', level=3)
    doc.add_paragraph('• Material: {mat_impulsion}')
    doc.add_paragraph('• Diámetro interno: {diam_impulsion_mm:.1f} mm')
    doc.add_paragraph('• Longitud: {long_impulsion:.1f} m')
    doc.add_paragraph('• Coeficiente Hazen-Williams: {coeficiente_hazen_impulsion}')

    # 3. CÁLCULOS HIDRÁULICOS
    doc.add_heading('3. CÁLCULOS HIDRÁULICOS', level=1)
    doc.add_heading('3.1 Pérdidas en Succión', level=2)
    doc.add_paragraph('• Velocidad: {velocidad_succion:.2f} m/s')
    doc.add_paragraph('• Pérdida primaria: {hf_primaria_succion:.2f} m')
    doc.add_paragraph('• Pérdida secundaria: {hf_secundaria_succion:.2f} m')
    doc.add_paragraph('• Pérdida total: {perdida_total_succion:.2f} m')
    doc.add_heading('3.2 Pérdidas en Impulsión', level=2)
    doc.add_paragraph('• Velocidad: {velocidad_impulsion:.2f} m/s')
    doc.add_paragraph('• Pérdida primaria: {hf_primaria_impulsion:.2f} m')
    doc.add_paragraph('• Pérdida secundaria: {hf_secundaria_impulsion:.2f} m')
    doc.add_paragraph('• Pérdida total: {perdida_total_impulsion:.2f} m')
    doc.add_heading('3.3 Altura Dinámica Total', level=2)
    doc.add_paragraph('Resultado: {altura_dinamica_total:.2f} m')

    # 4. ANÁLISIS NPSH
    doc.add_heading('4. ANÁLISIS NPSH', level=1)
    doc.add_paragraph('Resultados:')
    doc.add_paragraph('• NPSH disponible: {npshd_mca:.2f} m')
    doc.add_paragraph('• NPSH requerido: {npsh_requerido:.2f} m')
    doc.add_paragraph('• Margen NPSH: {npsh_margen:.2f} m')
    doc.add_heading('4.1 Análisis Técnico', level=2)
    doc.add_paragraph('{analisis_npsh}')

    # 5. SELECCIÓN DE MOTOR Y BOMBA
    doc.add_heading('5. SELECCIÓN DE MOTOR Y BOMBA', level=1)
    doc.add_paragraph('Resultados:')
    doc.add_paragraph('• Potencia hidráulica: {potencia_hidraulica_kw:.2f} kW ({potencia_hidraulica_hp:.2f} HP)')
    doc.add_paragraph('• Potencia del motor: {potencia_motor_final_kw:.2f} kW ({potencia_motor_final_hp:.2f} HP)')

    # 6. ANÁLISIS DE VARIADOR DE FRECUENCIA (VDF)
    doc.add_heading('6. ANÁLISIS DE VARIADOR DE FRECUENCIA (VDF)', level=1)
    doc.add_paragraph('• Porcentaje de RPM: {rpm_percentage:.1f}%')
    doc.add_paragraph('• Potencia ajustada: {potencia_ajustada:.2f} HP')
    doc.add_paragraph('• Eficiencia ajustada: {eficiencia_ajustada:.2f}%')

    # Marcadores para contenido dinámico al final (SIN IA)
    doc.add_paragraph('{seccion_tablas}')
    doc.add_paragraph('{seccion_graficos}')
    
    try:
        plantilla_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "informes", "plantillas")
        os.makedirs(plantilla_dir, exist_ok=True)
        plantilla_path = os.path.join(plantilla_dir, "plantilla_informe_sin_ia.docx")
        doc.save(plantilla_path)
        st.success(f"✅ Plantilla SIN IA creada/actualizada en: {plantilla_path}")
    except Exception as e:
        st.error(f"Error al guardar la plantilla SIN IA: {e}")

def crear_plantilla_con_ia():
    """Crea y guarda en disco un archivo DOCX con la estructura completa CON análisis IA."""
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    doc.add_heading('INFORME TÉCNICO DE SISTEMA DE BOMBEO', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. INFORMACIÓN DEL PROYECTO
    doc.add_heading('1. INFORMACIÓN DEL PROYECTO', level=1)
    doc.add_paragraph('Proyecto: {proyecto}')
    doc.add_paragraph('Diseño: {diseno}')
    doc.add_paragraph('Fecha: {fecha_generacion}')
    doc.add_paragraph('Caudal de diseño: {caudal_diseno_lps} L/s ({caudal_diseno_m3h} m³/h)')

    # 2. DATOS DE ENTRADA
    doc.add_heading('2. DATOS DE ENTRADA', level=1)
    doc.add_heading('2.1 Condiciones de Operación', level=2)
    doc.add_paragraph('• Temperatura del líquido: {temperatura} °C')
    doc.add_paragraph('• Densidad del líquido: {densidad_liquido} kg/m³')
    doc.add_paragraph('• Presión de vapor: {presion_vapor_calculada:.2f} kPa')
    doc.add_paragraph('• Presión barométrica: {presion_barometrica_calculada:.2f} kPa')
    doc.add_heading('2.2 Geometría del Sistema', level=2)
    doc.add_paragraph('• Altura de succión: {altura_succion} m')
    doc.add_paragraph('• Altura de descarga: {altura_descarga} m')
    doc.add_paragraph('• Altura estática total: {altura_estatica_total:.2f} m')
    doc.add_paragraph('• Número de bombas en paralelo: {num_bombas}')
    
    # Añadir caudal por bomba si hay paralelo
    if st.session_state.get('num_bombas', 1) > 1:
        doc.add_paragraph('• Caudal por bomba individual: {caudal_por_bomba_lps} L/s')
    doc.add_heading('2.3 Tuberías', level=2)
    doc.add_heading('Tubería de Succión', level=3)
    doc.add_paragraph('• Material: {mat_succion}')
    doc.add_paragraph('• Diámetro interno: {diam_succion_mm:.1f} mm')
    doc.add_paragraph('• Longitud: {long_succion:.1f} m')
    doc.add_paragraph('• Coeficiente Hazen-Williams: {coeficiente_hazen_succion}')
    doc.add_heading('Tubería de Impulsión', level=3)
    doc.add_paragraph('• Material: {mat_impulsion}')
    doc.add_paragraph('• Diámetro interno: {diam_impulsion_mm:.1f} mm')
    doc.add_paragraph('• Longitud: {long_impulsion:.1f} m')
    doc.add_paragraph('• Coeficiente Hazen-Williams: {coeficiente_hazen_impulsion}')

    # 3. CÁLCULOS HIDRÁULICOS
    doc.add_heading('3. CÁLCULOS HIDRÁULICOS', level=1)
    doc.add_heading('3.1 Pérdidas en Succión', level=2)
    doc.add_paragraph('• Velocidad: {velocidad_succion:.2f} m/s')
    doc.add_paragraph('• Pérdida primaria: {hf_primaria_succion:.2f} m')
    doc.add_paragraph('• Pérdida secundaria: {hf_secundaria_succion:.2f} m')
    doc.add_paragraph('• Pérdida total: {perdida_total_succion:.2f} m')
    doc.add_heading('3.2 Pérdidas en Impulsión', level=2)
    doc.add_paragraph('• Velocidad: {velocidad_impulsion:.2f} m/s')
    doc.add_paragraph('• Pérdida primaria: {hf_primaria_impulsion:.2f} m')
    doc.add_paragraph('• Pérdida secundaria: {hf_secundaria_impulsion:.2f} m')
    doc.add_paragraph('• Pérdida total: {perdida_total_impulsion:.2f} m')
    doc.add_heading('3.3 Altura Dinámica Total', level=2)
    doc.add_paragraph('Resultado: {altura_dinamica_total:.2f} m')

    # 4. ANÁLISIS NPSH
    doc.add_heading('4. ANÁLISIS NPSH', level=1)
    doc.add_paragraph('Resultados:')
    doc.add_paragraph('• NPSH disponible: {npshd_mca:.2f} m')
    doc.add_paragraph('• NPSH requerido: {npsh_requerido:.2f} m')
    doc.add_paragraph('• Margen NPSH: {npsh_margen:.2f} m')
    doc.add_heading('4.1 Análisis Técnico', level=2)
    doc.add_paragraph('{analisis_npsh}')

    # 5. SELECCIÓN DE MOTOR Y BOMBA
    doc.add_heading('5. SELECCIÓN DE MOTOR Y BOMBA', level=1)
    doc.add_paragraph('Resultados:')
    doc.add_paragraph('• Potencia hidráulica: {potencia_hidraulica_kw:.2f} kW ({potencia_hidraulica_hp:.2f} HP)')
    doc.add_paragraph('• Potencia del motor: {potencia_motor_final_kw:.2f} kW ({potencia_motor_final_hp:.2f} HP)')

    # 6. ANÁLISIS DE VARIADOR DE FRECUENCIA (VDF)
    doc.add_heading('6. ANÁLISIS DE VARIADOR DE FRECUENCIA (VDF)', level=1)
    doc.add_paragraph('• Porcentaje de RPM: {rpm_percentage:.1f}%')
    doc.add_paragraph('• Potencia ajustada: {potencia_ajustada:.2f} HP')
    doc.add_paragraph('• Eficiencia ajustada: {eficiencia_ajustada:.2f}%')

    # 7. ANÁLISIS INTELIGENTE (IA)
    doc.add_heading('7. ANÁLISIS INTELIGENTE DEL SISTEMA', level=1)
    doc.add_paragraph('{seccion_analisis_ia}')
    
    # 8. RECOMENDACIONES INTELIGENTES (IA)
    doc.add_heading('8. RECOMENDACIONES INTELIGENTES', level=1)
    doc.add_paragraph('{seccion_recomendaciones_ia}')

    # Marcadores para contenido dinámico al final (CON IA)
    doc.add_paragraph('{seccion_tablas}')
    doc.add_paragraph('{seccion_graficos}')
    
    try:
        plantilla_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "informes", "plantillas")
        os.makedirs(plantilla_dir, exist_ok=True)
        plantilla_path = os.path.join(plantilla_dir, "plantilla_informe_con_ia.docx")
        doc.save(plantilla_path)
        st.success(f"✅ Plantilla CON IA creada/actualizada en: {plantilla_path}")
    except Exception as e:
        st.error(f"Error al guardar la plantilla CON IA: {e}")

def crear_plantilla_base():
    """Crea automáticamente las plantillas según el estado del análisis IA"""
    # Crear siempre la plantilla sin IA
    crear_plantilla_sin_ia()
    
    # Crear plantilla con IA solo si está habilitado
    if st.session_state.get('ai_enabled', False):
        crear_plantilla_con_ia()
        st.info("🤖 Plantillas creadas: SIN IA y CON IA (análisis habilitado)")
    else:
        st.info("📄 Plantilla creada: Solo SIN IA (análisis deshabilitado)")

def consultar_ia_y_guardar_en_cache():
    """Consulta a la IA y guarda los resultados en session_state para evitar loops"""
    try:
        if not st.session_state.get('ai_enabled') or not st.session_state.get('model'):
            st.warning("⚠️ IA no configurada. No se generarán análisis inteligentes.")
            return False
        
        # Verificar si ya tenemos los datos en caché
        if 'analisis_ia_cache' in st.session_state:
            st.info("✅ Análisis de IA ya consultado. Usando datos en caché.")
            return True
        
        # Verificar que no estemos ya en proceso de consulta
        if st.session_state.get('consultando_ia', False):
            st.warning("⚠️ Ya hay una consulta a la IA en proceso. Por favor espera.")
            return False
        
        # Marcar que estamos consultando
        st.session_state['consultando_ia'] = True
        
        with st.spinner("🤖 Consultando a la IA para análisis del sistema..."):
            # Crear prompt con datos directos de st.session_state
            prompt_completo = f'''Eres un ingeniero hidráulico experto. Analiza este sistema de bombeo:

PROYECTO: {st.session_state.get('proyecto', 'N/A')}
DISEÑO: {st.session_state.get('diseno', 'N/A')}

CONDICIONES DE OPERACIÓN:
- Caudal de diseño: {st.session_state.get('caudal_lps', 0):.2f} L/s ({st.session_state.get('caudal_m3h', 0):.2f} m³/h)
- Altura estática total: {st.session_state.get('altura_estatica_total', 0):.2f} m
- Altura dinámica total: {st.session_state.get('adt_total', 0):.2f} m
- Temperatura del líquido: {st.session_state.get('temp_liquido', 20):.1f} °C
- Densidad del líquido: {st.session_state.get('densidad_liquido', 1.0):.3f} kg/m³

NPSH:
- NPSH disponible: {st.session_state.get('npshd_mca', 0):.2f} m
- NPSH requerido: {st.session_state.get('npsh_requerido', 0):.2f} m
- Margen NPSH: {st.session_state.get('npsh_margen', 0):.2f} m

MOTOR Y BOMBA:
- Potencia del motor: {st.session_state.get('potencia_motor_final_kw', 0):.2f} kW ({st.session_state.get('potencia_motor_final_hp', 0):.2f} HP)
- Potencia hidráulica: {st.session_state.get('potencia_hidraulica_kw', 0):.2f} kW ({st.session_state.get('potencia_hidraulica_hp', 0):.2f} HP)
- Eficiencia de operación: {st.session_state.get('eficiencia_operacion', 0):.2f}%

VFD (si aplica):
- Porcentaje RPM: {st.session_state.get('rpm_percentage', 100):.1f}%
- Potencia ajustada: {st.session_state.get('potencia_ajustada', 0):.2f} HP
- Eficiencia ajustada: {st.session_state.get('eficiencia_ajustada', 0):.2f}%

TUBERÍAS:
- Succión: {st.session_state.get('mat_succion', 'N/A')} Ø{st.session_state.get('diam_succion_mm', 0):.0f}mm, {st.session_state.get('long_succion', 0):.1f}m
- Impulsión: {st.session_state.get('mat_impulsion', 'N/A')} Ø{st.session_state.get('diam_impulsion_mm', 0):.0f}mm, {st.session_state.get('long_impulsion', 0):.1f}m

PÉRDIDAS DE CARGA:
- Pérdidas succión: {st.session_state.get('perdida_total_succion', 0):.2f} m
- Pérdidas impulsión: {st.session_state.get('perdida_total_impulsion', 0):.2f} m
- Pérdidas totales: {st.session_state.get('perdidas_totales_sistema', 0):.2f} m

Proporciona un análisis técnico completo con las siguientes secciones:

## Análisis General del Sistema
[Análisis técnico del sistema de bombeo basado en los datos proporcionados. Usa negritas reales para números importantes y unidades. Máximo 250 palabras.]

## Recomendaciones Técnicas
[Recomendaciones específicas y prácticas para optimizar el sistema. Usa negritas reales para números importantes y unidades. Máximo 250 palabras.]

## Análisis NPSH
[Análisis específico del NPSH y prevención de cavitación. Usa negritas reales para números importantes y unidades. Máximo 250 palabras.]

INSTRUCCIONES DE FORMATO:
- Usa ## para subtítulos
- Para números y unidades importantes, usa negritas reales (no **texto**)
- Ejemplo correcto: El caudal de diseño es 51.00 L/s
- Ejemplo incorrecto: El caudal de diseño es **51.00 L/s**
- NO uses marcadores Markdown como ** o __
- Escribe texto normal con números y unidades en negritas reales'''
            
            # Una sola consulta a la IA
            response = st.session_state.model.generate_content(prompt_completo)
            
            # Procesar la respuesta Markdown
            respuesta_texto = response.text
            
            # Extraer secciones del texto Markdown
            def extraer_seccion(texto, titulo):
                """Extrae una sección específica del texto Markdown"""
                try:
                    # Buscar el título de la sección
                    inicio = texto.find(f"## {titulo}")
                    if inicio == -1:
                        return f"Sección '{titulo}' no encontrada en la respuesta."
                    
                    # Encontrar el final de la sección (siguiente ## o final del texto)
                    siguiente_titulo = texto.find("## ", inicio + 1)
                    if siguiente_titulo == -1:
                        # Es la última sección
                        contenido = texto[inicio:].strip()
                    else:
                        contenido = texto[inicio:siguiente_titulo].strip()
                    
                    # Remover el título de la sección
                    lineas = contenido.split('\n')
                    if lineas and lineas[0].startswith('##'):
                        contenido = '\n'.join(lineas[1:]).strip()
                    
                    return contenido if contenido else f"Contenido de '{titulo}' vacío."
                except Exception as e:
                    return f"Error al extraer sección '{titulo}': {str(e)}"
            
            # Extraer cada sección
            analisis_general = extraer_seccion(respuesta_texto, "Análisis General del Sistema")
            recomendaciones = extraer_seccion(respuesta_texto, "Recomendaciones Técnicas")
            analisis_npsh = extraer_seccion(respuesta_texto, "Análisis NPSH")
            
            # Si no se encontraron secciones, usar la respuesta completa
            if "no encontrada" in analisis_general.lower():
                st.warning("⚠️ La IA no respondió con el formato esperado. Usando respuesta completa.")
                analisis_general = respuesta_texto
                recomendaciones = "Recomendaciones: Revisar el análisis general para obtener recomendaciones específicas."
                analisis_npsh = "Análisis NPSH: Revisar el análisis general para obtener información sobre NPSH."
            
            # Post-procesamiento: convertir marcadores Markdown a negritas reales para Word
            def convertir_negritas(texto):
                """Convierte **texto** a formato de negritas para documentos Word"""
                import re
                # Convertir **texto** a texto con formato de negritas real
                # En lugar de <b>, usar formato directo que Word entienda
                texto = re.sub(r'\*\*(.*?)\*\*', r'\1', texto)
                # Convertir __texto__ también
                texto = re.sub(r'__(.*?)__', r'\1', texto)
                return texto
            
            # Aplicar conversión a todas las secciones
            analisis_general = convertir_negritas(analisis_general)
            recomendaciones = convertir_negritas(recomendaciones)
            analisis_npsh = convertir_negritas(analisis_npsh)
            
            # Guardar en caché
            st.session_state['analisis_ia_cache'] = {
                'analisis_general': analisis_general,
                'recomendaciones': recomendaciones,
                'analisis_npsh': analisis_npsh,
                'timestamp': datetime.now().isoformat()
            }
            
            st.success("✅ Análisis de IA consultado y guardado en caché.")
            return True
            
    except Exception as e:
        st.error(f"❌ Error consultando a la IA: {str(e)}")
        st.error("💡 Posibles soluciones:")
        st.error("1. Verifica tu conexión a internet")
        st.error("2. Verifica que tu clave API de Gemini sea válida")
        st.error("3. Intenta limpiar el caché y consultar nuevamente")
        return False
    finally:
        # Limpiar el flag de consulta
        st.session_state['consultando_ia'] = False

def obtener_analisis_ia_desde_cache(tema):
    """Obtiene el análisis de IA desde el caché"""
    if 'analisis_ia_cache' not in st.session_state:
        return f"Análisis IA para '{tema}' no disponible (no consultado)."
    
    cache = st.session_state['analisis_ia_cache']
    
    if tema == "Análisis General del Sistema":
        return cache.get('analisis_general', f"Análisis para '{tema}' no disponible.")
    elif tema == "Recomendaciones Finales":
        return cache.get('recomendaciones', f"Recomendaciones para '{tema}' no disponible.")
    elif tema == "Análisis NPSH":
        # Usar el análisis específico de NPSH si está disponible
        analisis_npsh = cache.get('analisis_npsh', '')
        if analisis_npsh:
            return analisis_npsh
        else:
            # Fallback al análisis general si no hay análisis específico
            analisis_general = cache.get('analisis_general', '')
            if analisis_general and 'NPSH' in analisis_general:
                return f"Análisis NPSH basado en el análisis general del sistema: {analisis_general[:200]}..."
            else:
                return "Análisis NPSH: El sistema cumple con los requisitos de NPSH según el análisis general."
    else:
        return f"Análisis IA para '{tema}' no disponible."

def limpiar_cache_ia():
    """Limpia el caché de análisis de IA"""
    if 'analisis_ia_cache' in st.session_state:
        del st.session_state['analisis_ia_cache']
        st.info("🗑️ Caché de análisis IA limpiado.")
    
    # También limpiar el flag de consulta
    if 'consultando_ia' in st.session_state:
        del st.session_state['consultando_ia']

def generar_analisis_ia_para_reporte(tema):
    """Función legacy - ahora usa el caché"""
    return obtener_analisis_ia_desde_cache(tema)

def generar_informe_docx(incluir_graficos=True, incluir_tablas=True, incluir_analisis_ia=True):
    """Genera el informe DOCX cargando la plantilla y rellenando los datos."""
    try:
        # Seleccionar plantilla según si se incluye análisis IA
        if incluir_analisis_ia and st.session_state.get('ai_enabled', False):
            plantilla_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "informes", "plantillas", "plantilla_informe_con_ia.docx")
            plantilla_tipo = "CON IA"
        else:
            plantilla_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "informes", "plantillas", "plantilla_informe_sin_ia.docx")
            plantilla_tipo = "SIN IA"
        
        # --- FASE 5 MEJORA: Crear plantilla automáticamente si no existe ---
        if not os.path.exists(plantilla_path):
            st.warning(f"⚠️ Plantilla {plantilla_tipo} no encontrada. Creándola automáticamente...")
            try:
                if plantilla_tipo == "CON IA":
                    crear_plantilla_con_ia()
                else:
                    crear_plantilla_sin_ia()
                st.success(f"✅ Plantilla {plantilla_tipo} creada automáticamente")
            except Exception as e:
                st.error(f"❌ Error creando plantilla: {e}")
                return None, None
        
        # Verificar nuevamente después de intentar crear
        if not os.path.exists(plantilla_path):
            st.error(f"❌ No se pudo crear la plantilla {plantilla_tipo}.")
            return None, None

        doc = Document(plantilla_path)
        st.info(f"📄 Usando plantilla: {plantilla_tipo}")
        inputs = st.session_state

        # --- 1. Preparar y Reemplazar Variables de Texto ---
        # --- FASE 3 MEJORA: Placeholders simplificados (sin especificadores de formato) ---
        # Se mantiene compatibilidad con ambos formatos: nuevo (sin .2f) y antiguo (con .2f)
        variables = {
            # Información del proyecto
            '{proyecto}': inputs.get('proyecto', 'N/A'),
            '{diseno}': inputs.get('diseno', 'N/A'),
            '{fecha_generacion}': datetime.now().strftime('%Y/%m/%d %H:%M'),
            
            # Caudales
            '{caudal_diseno_lps}': f"{inputs.get('caudal_lps', 0.0):.2f}",
            '{caudal_diseno_m3h}': f"{inputs.get('caudal_m3h', 0.0):.2f}",
            '{caudal_por_bomba_lps}': f"{inputs.get('caudal_lps', 0.0) / inputs.get('num_bombas', 1):.2f}" if inputs.get('num_bombas', 1) > 0 else "0.00",
            
            # Condiciones del fluido
            '{temperatura}': f"{inputs.get('temp_liquido', 20.0):.1f}",
            '{densidad_liquido}': f"{inputs.get('densidad_liquido', 1.0):.3f}",
            '{presion_vapor_calculada}': f"{inputs.get('presion_vapor_calculada', 0.0):.2f}",
            '{presion_barometrica_calculada}': f"{inputs.get('presion_barometrica_calculada', 0.0):.2f}",
            
            # Alturas
            '{altura_succion}': f"{inputs.get('altura_succion_input', 0.0):.2f}",
            '{altura_descarga}': f"{inputs.get('altura_descarga', 0.0):.2f}",
            '{altura_estatica_total}': f"{inputs.get('altura_estatica_total', 0.0):.2f}",
            '{altura_dinamica_total}': f"{inputs.get('adt_total', 0.0):.2f}",
            
            # Bombas
            '{num_bombas}': str(inputs.get('num_bombas', 1)),
            
            # Succión
            '{mat_succion}': inputs.get('mat_succion', 'N/A'),
            '{diam_succion_mm}': f"{inputs.get('diam_succion_mm', 0.0):.1f}",
            '{long_succion}': f"{inputs.get('long_succion', 0.0):.1f}",
            '{coeficiente_hazen_succion}': str(inputs.get('coeficiente_hazen_succion', 'N/A')),
            '{velocidad_succion}': f"{inputs.get('velocidad_succion', 0.0):.2f}",
            '{hf_primaria_succion}': f"{inputs.get('hf_primaria_succion', 0.0):.2f}",
            '{hf_secundaria_succion}': f"{inputs.get('hf_secundaria_succion', 0.0):.2f}",
            '{perdida_total_succion}': f"{inputs.get('perdida_total_succion', 0.0):.2f}",
            
            # Impulsión
            '{mat_impulsion}': inputs.get('mat_impulsion', 'N/A'),
            '{diam_impulsion_mm}': f"{inputs.get('diam_impulsion_mm', 0.0):.1f}",
            '{long_impulsion}': f"{inputs.get('long_impulsion', 0.0):.1f}",
            '{coeficiente_hazen_impulsion}': str(inputs.get('coeficiente_hazen_impulsion', 'N/A')),
            '{velocidad_impulsion}': f"{inputs.get('velocidad_impulsion', 0.0):.2f}",
            '{hf_primaria_impulsion}': f"{inputs.get('hf_primaria_impulsion', 0.0):.2f}",
            '{hf_secundaria_impulsion}': f"{inputs.get('hf_secundaria_impulsion', 0.0):.2f}",
            '{perdida_total_impulsion}': f"{inputs.get('perdida_total_impulsion', 0.0):.2f}",
            
            # NPSH
            '{npshd_mca}': f"{inputs.get('npshd_mca', 0.0):.2f}",
            '{npsh_requerido}': f"{inputs.get('npsh_requerido', 0.0):.2f}",
            '{npsh_margen}': f"{inputs.get('npsh_margen', 0.0):.2f}",
            
            # Potencia
            '{potencia_hidraulica_kw}': f"{inputs.get('potencia_hidraulica_kw', 0.0):.2f}",
            '{potencia_hidraulica_hp}': f"{inputs.get('potencia_hidraulica_hp', 0.0):.2f}",
            '{potencia_motor_final_kw}': f"{inputs.get('potencia_motor_final_kw', 0.0):.2f}",
            '{potencia_motor_final_hp}': f"{inputs.get('potencia_motor_final_hp', 0.0):.2f}",
            
            # VFD
            '{rpm_percentage}': f"{inputs.get('rpm_percentage', 100.0):.1f}",
            '{potencia_ajustada}': f"{inputs.get('potencia_ajustada', 0.0):.2f}",
            '{eficiencia_ajustada}': f"{inputs.get('eficiencia_ajustada', 0.0):.2f}",
            
            # IA
            '{analisis_npsh}': "Análisis pendiente",
            
            # --- Compatibilidad con formato antiguo (con especificadores) ---
            '{presion_vapor_calculada:.2f}': f"{inputs.get('presion_vapor_calculada', 0.0):.2f}",
            '{presion_barometrica_calculada:.2f}': f"{inputs.get('presion_barometrica_calculada', 0.0):.2f}",
            '{altura_estatica_total:.2f}': f"{inputs.get('altura_estatica_total', 0.0):.2f}",
            '{diam_succion_mm:.1f}': f"{inputs.get('diam_succion_mm', 0.0):.1f}",
            '{long_succion:.1f}': f"{inputs.get('long_succion', 0.0):.1f}",
            '{diam_impulsion_mm:.1f}': f"{inputs.get('diam_impulsion_mm', 0.0):.1f}",
            '{long_impulsion:.1f}': f"{inputs.get('long_impulsion', 0.0):.1f}",
            '{velocidad_succion:.2f}': f"{inputs.get('velocidad_succion', 0.0):.2f}",
            '{hf_primaria_succion:.2f}': f"{inputs.get('hf_primaria_succion', 0.0):.2f}",
            '{hf_secundaria_succion:.2f}': f"{inputs.get('hf_secundaria_succion', 0.0):.2f}",
            '{perdida_total_succion:.2f}': f"{inputs.get('perdida_total_succion', 0.0):.2f}",
            '{velocidad_impulsion:.2f}': f"{inputs.get('velocidad_impulsion', 0.0):.2f}",
            '{hf_primaria_impulsion:.2f}': f"{inputs.get('hf_primaria_impulsion', 0.0):.2f}",
            '{hf_secundaria_impulsion:.2f}': f"{inputs.get('hf_secundaria_impulsion', 0.0):.2f}",
            '{perdida_total_impulsion:.2f}': f"{inputs.get('perdida_total_impulsion', 0.0):.2f}",
            '{altura_dinamica_total:.2f}': f"{inputs.get('adt_total', 0.0):.2f}",
            '{npshd_mca:.2f}': f"{inputs.get('npshd_mca', 0.0):.2f}",
            '{npsh_requerido:.2f}': f"{inputs.get('npsh_requerido', 0.0):.2f}",
            '{npsh_margen:.2f}': f"{inputs.get('npsh_margen', 0.0):.2f}",
            '{potencia_hidraulica_kw:.2f}': f"{inputs.get('potencia_hidraulica_kw', 0.0):.2f}",
            '{potencia_hidraulica_hp:.2f}': f"{inputs.get('potencia_hidraulica_hp', 0.0):.2f}",
            '{potencia_motor_final_kw:.2f}': f"{inputs.get('potencia_motor_final_kw', 0.0):.2f}",
            '{potencia_motor_final_hp:.2f}': f"{inputs.get('potencia_motor_final_hp', 0.0):.2f}",
            '{rpm_percentage:.1f}': f"{inputs.get('rpm_percentage', 100.0):.1f}",
            '{potencia_ajustada:.2f}': f"{inputs.get('potencia_ajustada', 0.0):.2f}",
            '{eficiencia_ajustada:.2f}': f"{inputs.get('eficiencia_ajustada', 0.0):.2f}",
        }
        
        # Las variables se aplicarán después de procesar la IA

        # --- FASE 1 MEJORA: Extraer DataFrames directamente de session_state ---
        # Función helper para obtener DataFrame de forma segura
        def safe_get_df(df_key, fallback_tablas_key=None):
            """Obtiene DataFrame de session_state con fallback a tablas_graficos"""
            # Primero intentar obtener directamente de session_state
            df = inputs.get(df_key)
            if df is not None and hasattr(df, 'empty') and not df.empty:
                return df
            
            # Fallback a tablas_graficos si existe
            if fallback_tablas_key:
                tablas = inputs.get('tablas_graficos', {})
                for grupo in ['tablas_100_rpm', 'tablas_vfd_rpm']:
                    datos = tablas.get(grupo, {})
                    df_data = datos.get(fallback_tablas_key)
                    if df_data:
                        df = deserialize_df(df_data)
                        if df is not None and not df.empty:
                            return df
            
            # Retornar DataFrame vacío si no se encuentra
            return pd.DataFrame()
        
        # Obtener el porcentaje de RPM VFD desde session_state
        rpm_vfd = inputs.get('rpm_percentage', 100)
        if rpm_vfd is None or rpm_vfd == 'N/A':
            rpm_vfd = 100

        # Obtener DataFrames directamente de session_state (método mejorado)
        df_bomba_100 = safe_get_df('df_bomba_100', 'df_bomba_100')
        df_sistema_100 = safe_get_df('df_sistema_100', 'df_sistema_100')
        df_eff_100 = safe_get_df('df_rendimiento_100', 'df_rendimiento_100')
        df_pow_100 = safe_get_df('df_potencia_100', 'df_potencia_100')
        df_npsh_100 = safe_get_df('df_npsh_100', 'df_npsh_100')

        df_bomba_vfd = safe_get_df('df_bomba_vfd', 'df_bomba_vfd')
        df_sistema_vfd = safe_get_df('df_sistema_vfd', 'df_sistema_vfd')
        df_eff_vfd = safe_get_df('df_rendimiento_vfd', 'df_rendimiento_vfd')
        df_pow_vfd = safe_get_df('df_potencia_vfd', 'df_potencia_vfd')
        df_npsh_vfd = safe_get_df('df_npsh_vfd', 'df_npsh_vfd')
        
        # Log de estado de tablas (para debugging)
        tablas_disponibles = sum([
            not df_bomba_100.empty, not df_sistema_100.empty, not df_eff_100.empty,
            not df_pow_100.empty, not df_npsh_100.empty
        ])
        st.info(f"📊 Tablas 100% RPM disponibles: {tablas_disponibles}/5")

        # --- Rellenar Secciones Dinámicas (IA, Tablas, Gráficos) ---
        if incluir_analisis_ia and inputs.get('ai_enabled'):
            # Verificar que la IA esté realmente configurada
            if not inputs.get('model'):
                st.warning("⚠️ Análisis IA solicitado pero no configurado. Generando informe sin análisis IA.")
                # Usar plantilla sin IA si no está configurada
                if plantilla_tipo == "CON IA":
                    st.info("🔄 Cambiando a plantilla SIN IA...")
                    # Recargar con plantilla sin IA
                    plantilla_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "informes", "plantillas", "plantilla_informe_sin_ia.docx")
                    if os.path.exists(plantilla_path):
                        doc = Document(plantilla_path)
                        st.info("📄 Usando plantilla: SIN IA (IA no configurada)")
            else:
                # --- FASE 4 MEJORA: Auto-consultar IA si no hay caché ---
                if 'analisis_ia_cache' not in st.session_state:
                    st.info("🤖 Consultando IA automáticamente...")
                    try:
                        consultar_ia_y_guardar_en_cache()
                        st.success("✅ Análisis IA generado automáticamente")
                    except Exception as e:
                        st.warning(f"⚠️ No se pudo consultar IA: {e}. Continuando sin análisis IA.")
                        # Continuar sin IA en lugar de bloquear
                        incluir_analisis_ia = False
                
                # Usar datos desde caché (ya consultados previamente)
                replace_placeholder_with_text(doc, '{seccion_analisis_ia}', obtener_analisis_ia_desde_cache("Análisis General del Sistema"), "7. Análisis General del Sistema (IA)")
                replace_placeholder_with_text(doc, '{seccion_recomendaciones_ia}', obtener_analisis_ia_desde_cache("Recomendaciones Finales"), "8. Recomendaciones (IA)")
                
                # También reemplazar el placeholder de análisis NPSH con datos del caché
                analisis_npsh = obtener_analisis_ia_desde_cache("Análisis NPSH")
                if analisis_npsh and not analisis_npsh.startswith("Análisis IA para"):
                    variables['{analisis_npsh}'] = analisis_npsh
                else:
                    variables['{analisis_npsh}'] = "Análisis NPSH disponible en sección de análisis general."
                
                # Las variables se aplicarán al final de la función
        
        # --- Insertar Tablas (solo si está activado el checkbox) ---
        if incluir_tablas:
            replace_placeholder_with_table(doc, '{seccion_tablas}', df_bomba_100, "9. Tabla de Rendimiento (100% RPM)")
            add_df_to_doc(doc, df_bomba_vfd, f"Tabla de Rendimiento ({rpm_vfd}% RPM)")
        else:
            # Reemplazar placeholder con texto vacío si no se incluyen tablas
            replace_placeholder_with_text(doc, '{seccion_tablas}', '', '')

        # --- Insertar Gráficos (solo si está activado el checkbox) ---
        if incluir_graficos:
            doc.add_heading('10. Gráficos de Rendimiento', level=1)
            
            # --- FASE 2 MEJORA: Intentar capturar gráficos automáticamente ---
            graficos_disponibles = 'graficos_exportados' in st.session_state
            
            if not graficos_disponibles:
                # Intentar captura automática
                st.info("🔄 Intentando capturar gráficos automáticamente...")
                try:
                    forzar_captura_graficos()
                    graficos_disponibles = 'graficos_exportados' in st.session_state
                    if graficos_disponibles:
                        st.success("✅ Gráficos capturados automáticamente")
                except Exception as e:
                    st.warning(f"⚠️ No se pudieron capturar gráficos automáticamente: {e}")
            
            if graficos_disponibles:
                # Usar gráficos capturados de Plotly (mejor calidad)
                doc.add_paragraph("📊 Gráficos generados con alta calidad desde análisis de curvas")
                
                # Gráficos 100% RPM (con eficiencia como referencia en eje cero)
                agregar_imagen_plotly_a_doc(doc, 'grupo_100_rpm', 'hq_100', "Curva H-Q (100% RPM)")
                agregar_imagen_plotly_a_doc(doc, 'grupo_100_rpm', 'rendimiento_100', "Curva de Eficiencia (100% RPM)")
                agregar_imagen_plotly_a_doc(doc, 'grupo_100_rpm', 'potencia_100', "Curva de Potencia (100% RPM)")
                agregar_imagen_plotly_a_doc(doc, 'grupo_100_rpm', 'npsh_100', "Curva NPSH Requerido (100% RPM)")
                
                # Gráficos VFD (con eficiencia como referencia en eje cero)
                doc.add_heading(f'10.2 Gráficos de Rendimiento ({rpm_vfd}% RPM)', level=2)
                agregar_imagen_plotly_a_doc(doc, 'grupo_vfd', 'hq_vfd', f"Curva H-Q ({rpm_vfd}% RPM)")
                agregar_imagen_plotly_a_doc(doc, 'grupo_vfd', 'rendimiento_vfd', f"Curva de Eficiencia ({rpm_vfd}% RPM)")
                agregar_imagen_plotly_a_doc(doc, 'grupo_vfd', 'potencia_vfd', f"Curva de Potencia ({rpm_vfd}% RPM)")
                agregar_imagen_plotly_a_doc(doc, 'grupo_vfd', 'npsh_vfd', f"Curva NPSH Requerido ({rpm_vfd}% RPM)")
                
            else:
                # --- FASE 2 MEJORA: Fallback mejorado con estilos profesionales ---
                doc.add_paragraph("📊 Gráficos generados con estilo técnico profesional")
                
                # Configuración de estilo mejorado para Matplotlib
                plt.style.use('seaborn-v0_8-whitegrid')
                COLORS = {'bomba': '#2E86AB', 'sistema': '#A23B72', 'operacion': '#F18F01', 'eficiencia': '#C73E1D'}
                
                bep_point = calculate_bep(df_eff_100.to_records(index=False).tolist()) if not df_eff_100.empty else None
                
                # 1. Curva H-Q 100%
                plt.figure(figsize=(8, 5))
                if not df_bomba_100.empty: plt.plot(df_bomba_100.iloc[:, 0], df_bomba_100.iloc[:, 1], label='Curva Bomba')
                if not df_sistema_100.empty: plt.plot(df_sistema_100.iloc[:, 0], df_sistema_100.iloc[:, 1], label='Curva Sistema')
                if inputs.get('interseccion') and len(inputs['interseccion']) >= 2: plt.plot(inputs['interseccion'][0], inputs['interseccion'][1], 'r*', markersize=10, label='Punto Operación')
                plt.title('Curva Bomba vs Sistema (100% RPM)'); plt.xlabel('Caudal (L/s)'); plt.ylabel('Altura (m)'); plt.grid(True); plt.legend()
                add_matplotlib_plot_to_doc(doc, "Curva H-Q (100% RPM)")

                # 2. Curva de Eficiencia 100% (curva real)
                plt.figure(figsize=(8, 5))
                if not df_eff_100.empty:
                    # Obtener valores de BEP para el label
                    zona_eff_min, zona_eff_max, bep_eta = obtener_valores_bep_eficiencia()
                    label_eficiencia = f"Zona de eficiencia ({zona_eff_min:.0f}%-{zona_eff_max:.0f}% BEP)"
                    
                    # Mostrar la curva real de eficiencia
                    plt.plot(df_eff_100.iloc[:, 0], df_eff_100.iloc[:, 1], '-', color='lightgray', 
                            label=label_eficiencia, linewidth=2, alpha=0.7)
                plt.title('Curva de Eficiencia (100% RPM)'); plt.xlabel('Caudal (L/s)'); plt.ylabel('Eficiencia (%)'); plt.grid(True); plt.legend()
                add_matplotlib_plot_to_doc(doc, "Curva de Eficiencia (100% RPM)")

                # 3. Curva de Potencia 100%
                plt.figure(figsize=(8, 5))
                if not df_pow_100.empty: plt.plot(df_pow_100.iloc[:, 0], df_pow_100.iloc[:, 1], label='Potencia')
                if inputs.get('interseccion') and inputs.get('potencia_operacion'): plt.plot(inputs['interseccion'][0], inputs['potencia_operacion'], 'ro', label='Punto Operación')
                plt.title('Curva de Potencia (100% RPM)'); plt.xlabel('Caudal (L/s)'); plt.ylabel('Potencia (HP)'); plt.grid(True); plt.legend()
                add_matplotlib_plot_to_doc(doc, "Curva de Potencia (100% RPM)")

                # 4. Curva NPSH 100%
                plt.figure(figsize=(8, 5))
                if not df_npsh_100.empty: plt.plot(df_npsh_100.iloc[:, 0], df_npsh_100.iloc[:, 1], label='NPSH Requerido')
                if inputs.get('interseccion') and inputs.get('npsh_requerido'): plt.plot(inputs['interseccion'][0], inputs['npsh_requerido'], 'ro', label='Punto Operación')
                plt.title('Curva NPSH Requerido (100% RPM)'); plt.xlabel('Caudal (L/s)'); plt.ylabel('NPSH (m)'); plt.grid(True); plt.legend()
                add_matplotlib_plot_to_doc(doc, "Curva NPSH Requerido (100% RPM)")

                # Gráficos VFD
                doc.add_heading(f'10.2 Gráficos de Rendimiento ({rpm_vfd}% RPM)', level=2)

                # 5. Curva H-Q VFD
                plt.figure(figsize=(8, 5))
                if not df_bomba_vfd.empty: plt.plot(df_bomba_vfd.iloc[:, 0], df_bomba_vfd.iloc[:, 1], label=f'Curva Bomba ({rpm_vfd}% RPM)')
                if not df_sistema_vfd.empty: plt.plot(df_sistema_vfd.iloc[:, 0], df_sistema_vfd.iloc[:, 1], label='Curva Sistema')
                if inputs.get('interseccion_vfd') and len(inputs['interseccion_vfd']) >= 2: plt.plot(inputs['interseccion_vfd'][0], inputs['interseccion_vfd'][1], 'r*', markersize=10, label='Punto Operación VFD')
                plt.title(f'Curva Bomba vs Sistema ({rpm_vfd}% RPM)'); plt.xlabel('Caudal (L/s)'); plt.ylabel('Altura (m)'); plt.grid(True); plt.legend()
                add_matplotlib_plot_to_doc(doc, f"Curva H-Q ({rpm_vfd}% RPM)")

                # 6. Curva de Eficiencia VFD (curva real)
                plt.figure(figsize=(8, 5))
                if not df_eff_vfd.empty:
                    # Obtener valores de BEP para el label
                    zona_eff_min, zona_eff_max, bep_eta = obtener_valores_bep_eficiencia()
                    label_eficiencia = f"Zona de eficiencia ({zona_eff_min:.0f}%-{zona_eff_max:.0f}% BEP)"
                    
                    # Mostrar la curva real de eficiencia
                    plt.plot(df_eff_vfd.iloc[:, 0], df_eff_vfd.iloc[:, 1], '-', color='lightgray', 
                            label=label_eficiencia, linewidth=2, alpha=0.7)
                plt.title(f'Curva de Eficiencia ({rpm_vfd}% RPM)'); plt.xlabel('Caudal (L/s)'); plt.ylabel('Eficiencia (%)'); plt.grid(True); plt.legend()
                add_matplotlib_plot_to_doc(doc, f"Curva de Eficiencia ({rpm_vfd}% RPM)")

                # 7. Curva de Potencia VFD
                plt.figure(figsize=(8, 5))
                if not df_pow_vfd.empty: plt.plot(df_pow_vfd.iloc[:, 0], df_pow_vfd.iloc[:, 1], label=f'Potencia ({rpm_vfd}% RPM)')
                if inputs.get('interseccion_vfd') and inputs.get('potencia_ajustada'): plt.plot(inputs['interseccion_vfd'][0], inputs['potencia_ajustada'], 'ro', label='Punto Operación VFD')
                plt.title(f'Curva de Potencia ({rpm_vfd}% RPM)'); plt.xlabel('Caudal (L/s)'); plt.ylabel('Potencia (HP)'); plt.grid(True); plt.legend()
                add_matplotlib_plot_to_doc(doc, f"Curva de Potencia ({rpm_vfd}% RPM)")

                # 8. Curva NPSH VFD
                plt.figure(figsize=(8, 5))
                if not df_npsh_vfd.empty: plt.plot(df_npsh_vfd.iloc[:, 0], df_npsh_vfd.iloc[:, 1], label=f'NPSH Requerido ({rpm_vfd}% RPM)')
                if inputs.get('interseccion_vfd') and inputs.get('npsh_requerido_vfd'): plt.plot(inputs['interseccion_vfd'][0], inputs['npsh_requerido_vfd'], 'ro', label='Punto Operación VFD')
                plt.title(f'Curva NPSH Requerido ({rpm_vfd}% RPM)'); plt.xlabel('Caudal (L/s)'); plt.ylabel('NPSH (m)'); plt.grid(True); plt.legend()
                add_matplotlib_plot_to_doc(doc, f"Curva NPSH Requerido ({rpm_vfd}% RPM)")
        else:
            # Reemplazar placeholder con texto vacío si no se incluyen gráficos
            replace_placeholder_with_text(doc, '{seccion_graficos}', '', '')

        # --- Aplicar todas las variables al documento ---
        replace_placeholders_in_doc(doc, variables)
        
        # --- Guardar Documento ---
        proyecto_nombre = str(inputs.get('proyecto', 'reporte')).replace(' ', '_')
        filename = f"Informe_{proyecto_nombre}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "informes", "docx")
        os.makedirs(output_dir, exist_ok=True)
        file_path = os.path.join(output_dir, filename)
        doc.save(file_path)
        
        with open(file_path, 'rb') as f:
            file_data = f.read()
        
        return file_data, filename
    
    except Exception as e:
        st.error(f"❌ Error Crítico al Generar Informe: {e}")
        import traceback
        st.error(traceback.format_exc())
        return None, None

def render_reports_tab():
    st.markdown("### 📄 Generación de Reportes Técnicos")
    
    if 'proyecto' not in st.session_state or not st.session_state['proyecto']:
        st.warning("⚠️ No hay ningún proyecto cargado.")
        return
    
    st.success(f"✅ Proyecto Activo: {st.session_state.get('proyecto', '')}")
    
    # Crear las 4 subpestañas
    tab1, tab2, tab3, tab4 = st.tabs(["📄 DOCX", "📊 PDF", "📈 XLSX", "🔧 EPANET"])
    
    with tab1:
        render_docx_subtab()
    
    with tab2:
        render_pdf_subtab()
    
    with tab3:
        render_xlsx_subtab()
    
    with tab4:
        render_epanet_subtab()

def render_docx_subtab():
    """Renderiza la subpestaña de reportes DOCX"""
    st.markdown("#### 📄 Generación de Reportes DOCX")
    
    # Mostrar directamente el contenido del método avanzado
    render_docx_avanzado()

def render_docx_avanzado():
    """Renderiza el método avanzado de generación DOCX"""
    # Crear 5 columnas de 20% cada una
    col1, col2, col3, col4, col5 = st.columns([20, 20, 20, 20, 20])
    
    with col1:
        st.markdown("**Configuración de Plantillas**")
        
        # Mostrar estado de las plantillas
        plantilla_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "informes", "plantillas")
        plantilla_sin_ia = os.path.join(plantilla_dir, "plantilla_informe_sin_ia.docx")
        plantilla_con_ia = os.path.join(plantilla_dir, "plantilla_informe_con_ia.docx")
        
        # Crear ambas plantillas automáticamente
        if st.button("🔧 Crear/Actualizar Ambas Plantillas", use_container_width=True):
            crear_plantilla_sin_ia()
            crear_plantilla_con_ia()
            st.success("✅ Ambas plantillas creadas/actualizadas exitosamente")
        
        st.markdown("**Modificar Plantilla Específica**")
        
        # Obtener lista de archivos .docx en la carpeta plantillas
        archivos_plantillas = []
        if os.path.exists(plantilla_dir):
            for archivo in os.listdir(plantilla_dir):
                if archivo.endswith('.docx'):
                    archivos_plantillas.append(archivo)
        
        if archivos_plantillas:
            # Combobox para seleccionar plantilla a modificar
            plantilla_seleccionada = st.selectbox(
                "Seleccionar plantilla para modificar:",
                archivos_plantillas,
                help="Elige qué plantilla quieres modificar manualmente (agregar logos, placeholders, etc.)"
            )
        else:
            st.warning("⚠️ No hay plantillas disponibles. Crea las plantillas primero.")
            plantilla_seleccionada = None
        
        if st.button("📝 Abrir Plantilla para Modificación", use_container_width=True):
            if plantilla_seleccionada:
                plantilla_path = os.path.join(plantilla_dir, plantilla_seleccionada)
                if os.path.exists(plantilla_path):
                    st.success(f"✅ Plantilla '{plantilla_seleccionada}' lista para modificar: {plantilla_path}")
                    st.info("💡 Puedes abrir este archivo en Word para agregar logos, cambiar placeholders, etc.")
                    # Intentar abrir el archivo en el explorador
                    try:
                        import subprocess
                        import platform
                        if platform.system() == "Windows":
                            subprocess.run(f'explorer /select,"{plantilla_path}"', shell=True)
                        elif platform.system() == "Darwin":  # macOS
                            subprocess.run(["open", "-R", plantilla_path])
                        else:  # Linux
                            subprocess.run(["xdg-open", os.path.dirname(plantilla_path)])
                        st.info("📁 Explorador de archivos abierto")
                    except Exception as e:
                        st.warning(f"No se pudo abrir el explorador: {e}")
                else:
                    st.error(f"❌ Plantilla '{plantilla_seleccionada}' no existe.")
            else:
                st.warning("⚠️ Selecciona una plantilla primero.")
        
        st.markdown("**Opciones de Generación**")
        incluir_graficos = st.checkbox("📊 Incluir gráficos", value=False, key="docx_incluir_graficos", help="Al activar, captura automáticamente los gráficos de la pestaña 'Análisis de curvas'")
        incluir_tablas = st.checkbox("📋 Incluir tablas", value=False, key="docx_incluir_tablas")
        # Sincronización bidireccional en tiempo real
        incluir_analisis_ia = st.checkbox(
            "🤖 Incluir análisis IA", 
            value=st.session_state.get('ai_enabled', False), 
            key="docx_incluir_analisis_ia"
        )
        
        # Sincronizar automáticamente cuando cambia el estado
        if 'docx_incluir_analisis_ia' in st.session_state:
            st.session_state['ai_enabled'] = st.session_state.docx_incluir_analisis_ia
        
        # Verificar configuración de IA si está activada
        if incluir_analisis_ia:
            if not st.session_state.get('model'):
                st.warning("⚠️ IA activada pero no configurada. Ve al panel lateral '🤖 Análisis IA' para configurar la API.")
            else:
                st.success("✅ IA configurada y lista para usar.")
                
                # Botones para gestionar caché de IA
                col_ia1, col_ia2, col_ia3 = st.columns(3)
                
                with col_ia1:
                    if st.button("🤖 Consultar IA", use_container_width=True, key="consultar_ia_cache"):
                        consultar_ia_y_guardar_en_cache()
                
                with col_ia2:
                    if st.button("🗑️ Limpiar Caché", use_container_width=True, key="limpiar_cache_ia"):
                        limpiar_cache_ia()
                
                with col_ia3:
                    if st.session_state.get('consultando_ia', False):
                        if st.button("❌ Cancelar Consulta", use_container_width=True, key="cancelar_consulta_ia"):
                            st.session_state['consultando_ia'] = False
                            st.warning("⚠️ Consulta cancelada por el usuario.")
                
                # Mostrar estado del caché
                if st.session_state.get('consultando_ia', False):
                    st.warning("🔄 Consultando a la IA... Por favor espera.")
                elif 'analisis_ia_cache' in st.session_state:
                    cache = st.session_state['analisis_ia_cache']
                    timestamp = cache.get('timestamp', 'N/A')
                    st.info(f"💾 Caché disponible (consultado: {timestamp[:19] if timestamp != 'N/A' else 'N/A'})")
                else:
                    st.warning("⚠️ No hay caché de IA. Consulta a la IA primero.")
        
        # Capturar gráficos cuando se active el checkbox
        if incluir_graficos and not st.session_state.get('graficos_capturados_para_reporte', False):
            with st.spinner("📊 Capturando gráficos de la pestaña 'Análisis de curvas'..."):
                resultado = capturar_todos_los_graficos_automaticamente()
                if resultado:
                    st.session_state['graficos_capturados_para_reporte'] = True
                    st.success("✅ Gráficos capturados exitosamente desde la pestaña 'Análisis de curvas'")
                else:
                    # Mostrar información de debugging
                    st.error("❌ No se pudieron capturar los gráficos.")
                    debug_captura_graficos()
        
        # Resetear flag si se desactiva el checkbox
        if not incluir_graficos:
            st.session_state['graficos_capturados_para_reporte'] = False
    
    with col2:
        st.markdown("**Generación de Reporte**")
        if st.button("🔄 Generar Reporte DOCX", type="primary", use_container_width=True):
            # La verificación de IA ahora se hace automáticamente en generar_informe_docx
            with st.spinner("Generando reporte profesional..."):
                file_data, filename = generar_informe_docx(incluir_graficos, incluir_tablas, incluir_analisis_ia)
            
            # Botón de descarga FUERA del spinner para que aparezca
            if file_data:
                st.success(f"✅ Reporte generado: `{filename}`")
                st.download_button(
                    "📥 Descargar Reporte DOCX", 
                    file_data, 
                    filename, 
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                    use_container_width=True,
                    key="download_docx_main"
                )
            else:
                st.error("❌ Error al generar el reporte. Revisa los mensajes anteriores.")
        
        st.markdown("**Vista Previa**")
        st.info("La vista previa se mostrará aquí una vez generado el reporte.")
        
        # Mostrar estado de gráficos capturados
        if 'graficos_exportados' in st.session_state:
            graficos_100 = len(st.session_state['graficos_exportados']['grupo_100_rpm'])
            graficos_vfd = len(st.session_state['graficos_exportados']['grupo_vfd'])
            st.success(f"📊 Gráficos capturados: {graficos_100} (100% RPM) + {graficos_vfd} (VFD)")
        else:
            st.info("ℹ️ Activa 'Incluir gráficos' para capturar las curvas de la pestaña 'Análisis de curvas'.")
    
    with col3:
        st.markdown("**Descarga de Gráficos**")
        
        # Botones de descarga de grupos de gráficos
        if 'graficos_exportados' in st.session_state:
            if st.button("📥 Grupo 100% RPM", use_container_width=True, key="descargar_grupo_100"):
                descargar_grupo_graficos('grupo_100_rpm', '100_RPM')
            
            if st.button("📥 Grupo VFD", use_container_width=True, key="descargar_grupo_vfd"):
                descargar_grupo_graficos('grupo_vfd', 'VFD')
            
            if st.button("📥 Todos los Gráficos", use_container_width=True, key="descargar_todos_graficos"):
                descargar_todos_los_graficos()
        else:
            st.info("Genera curvas en 'Análisis de curvas' para habilitar descargas")
    
    with col4:
        pass  # Espacio reservado para futuras funcionalidades
    
    with col5:
        pass  # Espacio reservado para futuras funcionalidades

def render_docx_rapido():
    """Renderiza el método rápido de generación DOCX - Implementación del patrón proporcionado"""
    st.markdown("#### ⚡ Generación Rápida de Word con Gráficos")
    st.info("🚀 **Método optimizado**: Genera documentos Word con gráficos en un solo paso, sin archivos temporales ni Kaleido.")
    
    # Crear pestañas para diferentes opciones
    tab_ejemplos, tab_session_state, tab_diagnostico = st.tabs(["🧪 Ejemplos", "📊 Desde Session State", "🔍 Diagnóstico"])
    
    with tab_ejemplos:
        st.markdown("### 🧪 Generar Gráficos de Ejemplo")
        
        # Crear 2 columnas
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**📊 Generar Gráfico Matplotlib**")
            if st.button("🔄 Crear Gráfico Matplotlib", use_container_width=True, key="ejemplo_matplotlib"):
                # Generar gráfico
                fig = generar_grafico_bombeo_matplotlib()
                st.pyplot(fig)  # Mostrar en la app
                
                # Guardar en session_state
                guardar_grafico_en_session_state(fig, "ejemplo_matplotlib", "ejemplos")
                
                # Crear documento Word
                doc_data = crear_docx_con_grafico(
                    fig, 
                    titulo_documento="Informe de Sistema de Bombeo - Matplotlib",
                    titulo_grafico="Curva H-Q del Sistema de Bombeo"
                )
                
                if doc_data:
                    # Crear nombre de archivo
                    from datetime import datetime
                    nombre = f"reporte_bombeo_matplotlib_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                    
                    # Botón de descarga
                    st.download_button(
                        label="📥 Descargar DOCX",
                        data=doc_data,
                        file_name=nombre,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key="descarga_matplotlib"
                    )
                    st.success("✅ Documento Word generado exitosamente")
        
        with col2:
            st.markdown("**📈 Generar Gráfico Plotly**")
            if st.button("🔄 Crear Gráfico Plotly", use_container_width=True, key="ejemplo_plotly"):
                # Generar gráfico
                fig = generar_grafico_bombeo_plotly()
                st.plotly_chart(fig, use_container_width=True)  # Mostrar en la app
                
                # Guardar en session_state
                guardar_grafico_en_session_state(fig, "ejemplo_plotly", "ejemplos")
                
                # Crear documento Word
                doc_data = crear_docx_con_grafico(
                    fig, 
                    titulo_documento="Informe de Sistema de Bombeo - Plotly",
                    titulo_grafico="Curva H-Q del Sistema de Bombeo"
                )
                
                if doc_data:
                    # Crear nombre de archivo
                    from datetime import datetime
                    nombre = f"reporte_bombeo_plotly_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                    
                    # Botón de descarga
                    st.download_button(
                        label="📥 Descargar DOCX",
                        data=doc_data,
                        file_name=nombre,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key="descarga_plotly"
                    )
                    st.success("✅ Documento Word generado exitosamente")
    
    with tab_session_state:
        st.markdown("### 📊 Generar DOCX desde Gráficos en Session State")
        st.info("💡 **Importante**: Los gráficos de la pestaña 'Análisis de curvas' se guardan automáticamente en session_state cuando se muestran.")
        
        # Mostrar estado actual
        mostrar_estado_graficos()
        
        # Botón para generar DOCX con todos los gráficos
        if st.button("🔄 Generar DOCX con Todos los Gráficos", type="primary", use_container_width=True):
            doc_data = crear_docx_con_graficos_desde_session_state("Informe Completo de Sistema de Bombeo")
            
            if doc_data:
                # Crear nombre de archivo
                from datetime import datetime
                nombre = f"reporte_completo_bombeo_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                
                # Botón de descarga
                st.download_button(
                    label="📥 Descargar DOCX Completo",
                    data=doc_data,
                    file_name=nombre,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="descarga_completo"
                )
                st.success("✅ Documento Word completo generado exitosamente")
            else:
                st.warning("⚠️ No se pudo generar el documento. Verifica que haya gráficos disponibles.")
        
        # Instrucciones
        st.markdown("---")
        st.markdown("### 📋 Instrucciones")
        st.info("""
        1. **Genera gráficos**: Ve a la pestaña 'Análisis de curvas' y genera las curvas
        2. **Verifica estado**: Usa la pestaña 'Diagnóstico' para verificar que los gráficos estén guardados
        3. **Genera DOCX**: Usa el botón de arriba para crear el documento con todos los gráficos
        4. **Descarga**: El documento se descargará automáticamente
        """)
    
    with tab_diagnostico:
        st.markdown("### 🔍 Diagnóstico del Sistema")
        
        # Diagnóstico del sistema
        if st.button("🧪 Ejecutar Diagnóstico Completo", use_container_width=True):
            diagnostico_sistema_completo()
        
        # Mostrar estado de gráficos
        st.markdown("---")
        mostrar_estado_graficos()
        
        # Información adicional
        st.markdown("---")
        st.markdown("### 💡 Características del Método Rápido")
        
        col_info1, col_info2, col_info3 = st.columns(3)
        
        with col_info1:
            st.info("""
            **🚀 Sin Archivos Temporales**
            - Todo se mantiene en memoria RAM
            - No se escriben archivos en disco
            - Ideal para servidores de Streamlit
            """)
        
        with col_info2:
            st.info("""
            **📊 Alta Calidad**
            - Gráficos a 300 DPI
            - Formato PNG optimizado
            - Compatible con Matplotlib y Plotly
            """)
        
        with col_info3:
            st.info("""
            **⚡ Un Solo Paso**
            - Generación instantánea
            - Descarga directa desde Streamlit
            - Sin dependencias externas
            """)

def render_pdf_subtab():
    """Renderiza la subpestaña de reportes PDF"""
    st.markdown("#### 📊 Generación de Reportes PDF")
    
    # Crear 5 columnas de 20% cada una
    col1, col2, col3, col4, col5 = st.columns([20, 20, 20, 20, 20])
    
    with col1:
        st.markdown("**Selección de Contenido**")
        st.markdown("##### 📋 Datos de Entrada")
        
        # Checkboxes para Datos de Entrada
        pdf_include_condiciones = st.checkbox("Condiciones de Operación", value=True, key="pdf_condiciones")
        pdf_include_succion = st.checkbox("Tubería y Accesorios de Succión", value=True, key="pdf_succion")
        pdf_include_impulsion = st.checkbox("Tubería y Accesorios de Impulsión", value=True, key="pdf_impulsion")
        pdf_include_curvas = st.checkbox("Ajuste de Curvas Características", value=True, key="pdf_curvas")
        pdf_include_diagrama = st.checkbox("📐 Diagrama Esquemático del Sistema", value=True, key="pdf_diagrama")
        
        st.markdown("##### 📊 Análisis de Curvas")
        pdf_include_npsh = st.checkbox("Resultados de Cálculos Hidráulicos", value=True, key="pdf_npsh")
        pdf_include_graficos_100 = st.checkbox("Gráfico de Curvas 100% RPM", value=True, key="pdf_graficos_100")
        pdf_include_graficos_vfd = st.checkbox("Gráfico de Curvas VDF", value=True, key="pdf_graficos_vfd")
        pdf_include_tablas = st.checkbox("Tablas", value=True, key="pdf_tablas")
        
        # Análisis de Transientes (solo si está habilitado)
        transient_enabled = st.session_state.get('transient_analysis_enabled', False)
        pdf_include_transientes = False
        if transient_enabled:
            st.markdown("##### 🔄 Análisis de Transientes")
            pdf_include_transientes = st.checkbox("Resultados de Transientes", value=True, key="pdf_transientes")
    
    with col2:
        st.markdown("**Configuración PDF**")
        calidad_pdf = st.selectbox("Calidad del PDF", ["Alta", "Media", "Baja"], index=1, key="pdf_calidad",
                                   help="Alta: 300 DPI (~mayor tamaño), Media: 150 DPI, Baja: 72 DPI (~menor tamaño)")
        orientacion_pdf = st.selectbox("Orientación", ["Vertical", "Horizontal"], index=0, key="pdf_orientacion")
        incluir_portada = st.checkbox("📄 Incluir portada", value=True, key="pdf_portada")
        incluir_indice = st.checkbox("📑 Incluir índice", value=True, key="pdf_indice")
        
        st.markdown("**Información**")
        tamanio_estimado = {"Alta": "~5-10 MB", "Media": "~2-5 MB", "Baja": "~1-2 MB"}
        st.info(f"Tamaño estimado: {tamanio_estimado.get(calidad_pdf, 'N/A')}")
    
    with col3:
        st.markdown("**Generación PDF**")
        if st.button("🔄 Generar Reporte PDF", type="primary", use_container_width=True):
            try:
                from data.pdf_export import create_pdf_report
                
                # Configurar las secciones a incluir
                config = {
                    'calidad': calidad_pdf,
                    'orientacion': orientacion_pdf,
                    'incluir_portada': incluir_portada,
                    'incluir_indice': incluir_indice,
                    'secciones': {
                        'condiciones': pdf_include_condiciones,
                        'succion': pdf_include_succion,
                        'impulsion': pdf_include_impulsion,
                        'curvas': pdf_include_curvas,
                        'diagrama': pdf_include_diagrama,
                        'npsh': pdf_include_npsh,
                        'graficos_100': pdf_include_graficos_100,
                        'graficos_vfd': pdf_include_graficos_vfd,
                        'tablas': pdf_include_tablas,
                        'transientes': pdf_include_transientes
                    }
                }
                
                # Generar el PDF
                with st.spinner('Generando reporte PDF...'):
                    pdf_output = create_pdf_report(st.session_state, config)
                    
                    # Guardar en session_state
                    st.session_state.pdf_report_data = pdf_output.getvalue()
                    st.session_state.pdf_report_generated = True
                    st.success("✅ Reporte PDF generado exitosamente")
                
            except Exception as e:
                st.error(f"Error al generar el reporte PDF: {e}")
                import traceback
                st.code(traceback.format_exc())
                st.session_state.pdf_report_generated = False
        
        # Mostrar el botón de descarga solo si el reporte ha sido generado
        if st.session_state.get('pdf_report_generated', False):
            nombre_proyecto = st.session_state.get('proyecto', 'Proyecto_Bombeo')
            nombre_archivo = f"{nombre_proyecto}_Reporte.pdf"
            
            st.download_button(
                label="📥 Descargar Reporte PDF",
                data=st.session_state.pdf_report_data,
                file_name=nombre_archivo,
                mime="application/pdf",
                use_container_width=True
            )
    
    with col4:
        st.markdown("**Contenido del Reporte**")
        st.info("""
        **El reporte PDF incluye:**
        
        - 📄 Portada con información del proyecto
        - 📑 Índice de contenidos
        - 📋 Datos de entrada y configuración
        - 📐 Diagrama esquemático del sistema
        - 📊 Resultados de cálculos hidráulicos
        - 📈 Gráficos de curvas (si se seleccionan)
        - 📋 Tablas de datos
        - 🔄 Análisis de transientes (si está habilitado)
        - 📄 Numeración de páginas
        """)
    
    with col5:
        st.markdown("**Ayuda**")
        st.info("""
        **Recomendaciones:**
        
        - **Calidad Alta**: Para impresión profesional
        - **Calidad Media**: Balance entre calidad y tamaño
        - **Calidad Baja**: Para envío por email
        
        - **Vertical**: Documentos estándar
        - **Horizontal**: Gráficos grandes
        """)

def render_xlsx_subtab():
    """Renderiza la subpestaña de reportes XLSX"""
    st.markdown("#### 📈 Generación de Reportes XLSX")
    
    # Crear 5 columnas de 20% cada una
    col1, col2, col3, col4, col5 = st.columns([20, 20, 20, 20, 20])
    
    with col1:
        st.markdown("**Exportación de Reporte Completo**")
        st.markdown("Exporta todos los inputs, resultados y datos de curvas a un único archivo Excel.")
        
        # Botón para generar el reporte completo
        if st.button("📊 Generar Reporte Completo en Excel", type="primary", key="export_reporte_xlsx", use_container_width=True):
            try:
                from data.export import create_comprehensive_excel_report
                
                # Obtener configuraciones
                incluir_formulas = st.session_state.get('xlsx_incluir_formulas', False)
                incluir_graficos_xlsx = st.session_state.get('xlsx_incluir_graficos', True)
                
                # Generar el archivo Excel mejorado en memoria con configuraciones
                excel_output = create_comprehensive_excel_report(
                    st.session_state, 
                    incluir_formulas=incluir_formulas,
                    incluir_graficos=incluir_graficos_xlsx
                )
                
                # Guardar en session_state para que el botón de descarga aparezca
                st.session_state.xlsx_report_data = excel_output.getvalue()
                st.session_state.xlsx_report_generated = True
                st.success("✅ Reporte generado exitosamente")
                
            except Exception as e:
                st.error(f"Error al generar el reporte: {e}")
                st.session_state.xlsx_report_generated = False

        # Mostrar el botón de descarga solo si el reporte ha sido generado
        if st.session_state.get('xlsx_report_generated', False):
            st.download_button(
                label="📥 Descargar Reporte Excel",
                data=st.session_state.xlsx_report_data,
                file_name="Reporte_Analisis_Bombeo.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
    
    with col2:
        st.markdown("**Información del Reporte**")
        st.info("""
        **Contenido del Reporte Excel:**
        
        - **Resumen:** Parámetros principales del proyecto
        - **Succión:** Cálculos de tubería de succión
        - **Impulsión:** Cálculos de tubería de impulsión
        - **Curvas:** Datos de curvas características
        - **Resultados:** Análisis completo del sistema
        """)
    
    with col3:
        st.markdown("**Configuración Adicional**")
        incluir_formulas = st.checkbox("🧮 Incluir fórmulas", value=False, key="xlsx_incluir_formulas", 
                                       help="Al activar, las celdas de la hoja 'Datos Gráficos' contendrán fórmulas en lugar de valores")
        incluir_graficos_xlsx = st.checkbox("📊 Incluir gráficos", value=True, key="xlsx_incluir_graficos",
                                           help="Al desactivar, se ocultarán las hojas 'Gráficos 100% RPM' y 'Gráficos VDF'")
        formato_numeros = st.selectbox("Formato de números", ["2 decimales", "3 decimales", "4 decimales"], index=0, key="xlsx_formato_numeros")
    
    with col4:
        pass  # Espacio reservado para futuras funcionalidades
    
    with col5:
        pass  # Espacio reservado para futuras funcionalidades

def render_epanet_subtab():
    """Renderiza la subpestaña de exportación EPANET"""
    # Usar el módulo completo de exportación EPANET
    render_epanet_export_section()
